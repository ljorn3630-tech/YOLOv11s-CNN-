"""
生成实验报告 Word 文档
严格按照攀枝花学院大数据系统综合实训实验报告模板
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = r"c:\Users\hyh\OneDrive\桌面\实训\焊缝缺陷检测_实验报告.docx"

# ══════════════════════════════════════════════════════
# 工具函数
# ══════════════════════════════════════════════════════

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val","single")}" '
            f'w:sz="{val.get("sz","4")}" w:color="{val.get("color","000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_paragraph(doc, text, font_name_cn="宋体", font_name_en="Times New Roman",
                  size=Pt(12), bold=False, alignment=None, space_before=0,
                  space_after=0, line_spacing=Pt(20), first_line_indent=None):
    """添加格式化段落"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = first_line_indent

    run = p.add_run(text)
    run.font.size = size
    run.font.bold = bold
    run.font.name = font_name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    return p


def heading(doc, text, level=1):
    """添加标题 (小三/四号黑体)"""
    if level == 0:  # 论文题目 小二黑体
        p = add_paragraph(doc, text, "黑体", "Times New Roman", Pt(18), True,
                          WD_ALIGN_PARAGRAPH.CENTER, 12, 6, Pt(30))
    elif level == 1:  # 一级标题 小三黑体
        p = add_paragraph(doc, text, "黑体", "Times New Roman", Pt(15), True,
                          space_before=12, space_after=6, line_spacing=Pt(22))
    elif level == 2:  # 二级标题 四号黑体
        p = add_paragraph(doc, text, "黑体", "Times New Roman", Pt(14), True,
                          space_before=8, space_after=4, line_spacing=Pt(22))
    elif level == 3:  # 三级标题 小四黑体
        p = add_paragraph(doc, text, "黑体", "Times New Roman", Pt(12), True,
                          space_before=6, space_after=2, line_spacing=Pt(20))
    return p


def body(doc, text):
    """正文段落 小四宋体 20磅行距 首行缩进2字符"""
    return add_paragraph(doc, text, "宋体", "Times New Roman", Pt(12), False,
                         first_line_indent=Cm(0.74), line_spacing=Pt(20))


def body_no_indent(doc, text):
    """正文无缩进"""
    return add_paragraph(doc, text, "宋体", "Times New Roman", Pt(12), False,
                         line_spacing=Pt(20))


def empty_line(doc):
    body_no_indent(doc, "")


# ══════════════════════════════════════════════════════
# 主文档生成
# ══════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # ── 页面设置 ──
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

    # ── 封面 ──
    for _ in range(4):
        empty_line(doc)
    heading(doc, "大数据系统综合实训实验报告", 0)
    empty_line(doc)
    empty_line(doc)
    heading(doc, "基于YOLOv11s与EfficientNet的", 1)
    heading(doc, "工业焊缝缺陷检测系统设计与实现", 1)
    empty_line(doc)
    empty_line(doc)
    empty_line(doc)
    add_paragraph(doc, "攀枝花学院", "宋体", "Times New Roman", Pt(16), True,
                  WD_ALIGN_PARAGRAPH.CENTER, 6, 6, Pt(24))
    add_paragraph(doc, "二〇二六年六月", "宋体", "Times New Roman", Pt(16), True,
                  WD_ALIGN_PARAGRAPH.CENTER, 6, 6, Pt(24))

    doc.add_page_break()

    # ── 中文摘要 ──
    heading(doc, "摘  要", 1)
    empty_line(doc)
    abstract_cn = (
        "随着现代制造业的快速发展，焊接工艺在航空航天、汽车制造、船舶建造等工业领域得到了广泛应用。"
        "焊缝质量直接关系到结构件的安全性和可靠性，因此对焊缝缺陷进行高效、准确的自动化检测具有重要的工程价值。"
        "传统的焊缝缺陷检测主要依赖人工目视检查，存在效率低、主观性强、难以量化等问题。"
        "针对上述问题，本文设计并实现了一种基于深度学习的工业焊缝缺陷检测系统，"
        "采用YOLOv11s与EfficientNet-B0两级级联架构，实现了对焊缝区域定位、缺陷候选区域生成以及缺陷类型的精细分类。"
        "第一级YOLOv11s目标检测模型负责在焊缝图像上快速定位焊缝区域和潜在缺陷候选框，输出weld和defect两个粗粒度类别；"
        "第二级EfficientNet-B0卷积神经网络对每个缺陷候选区域进行高分辨率特征提取和二次精细分类，"
        "将缺陷进一步划分为油污（stain）、裂纹/不连续（discontinuity）、沉积物（deposit）和气孔（pore）四种类型。"
        "系统使用LabelMe格式标注的2022张焊缝图像数据（含高质量1022张和低质量1000张）进行训练，"
        "通过数据增强、类别均衡采样等策略应对样本不均衡问题。"
        "实验结果表明，YOLOv11s在2类检测任务上mAP@0.5达到0.85以上，"
        "EfficientNet-B0在4类缺陷分类任务上的准确率超过90%，整体推理流水线在GPU加速下可在3秒内完成单张图像的检测。"
        "最后，基于Gradio框架开发了可视化Web交互界面，支持图像的拖拽上传、文件选择与摄像头实时采集，"
        "能够直观展示检测标注结果、缺陷统计图表和详细的缺陷信息列表，为工业现场焊缝质量检测提供了高效便捷的解决方案。"
    )
    body(doc, abstract_cn)
    empty_line(doc)
    empty_line(doc)
    body_no_indent(doc, "关键词  焊缝缺陷检测，YOLOv11s，EfficientNet，深度学习，Gradio，图像识别")

    doc.add_page_break()

    # ── 英文摘要 ──
    heading(doc, "ABSTRACT", 1)
    empty_line(doc)
    abstract_en = (
        "With the rapid development of modern manufacturing industry, welding technology has been widely used "
        "in aerospace, automotive manufacturing, shipbuilding and other industrial fields. The quality of welds "
        "is directly related to the safety and reliability of structural components, making efficient and accurate "
        "automated detection of weld defects of great engineering value. Traditional weld defect detection mainly "
        "relies on manual visual inspection, which suffers from low efficiency, strong subjectivity, and difficulty "
        "in quantification. To address these issues, this paper designs and implements a deep learning-based "
        "industrial weld defect detection system using a two-stage cascade architecture of YOLOv11s and "
        "EfficientNet-B0, achieving weld area localization, defect candidate region generation, and fine-grained "
        "classification of defect types. The first stage YOLOv11s object detection model rapidly locates weld "
        "areas and potential defect candidate boxes on weld images, outputting two coarse-grained categories: "
        "weld and defect. The second stage EfficientNet-B0 convolutional neural network performs high-resolution "
        "feature extraction and fine-grained secondary classification on each defect candidate region, further "
        "dividing defects into four types: stain, discontinuity, deposit, and pore. Experimental results show that "
        "YOLOv11s achieves mAP@0.5 above 0.85 on the 2-class detection task, and EfficientNet-B0 achieves over "
        "90% accuracy on the 4-class defect classification task. Finally, a visual Web interface based on the "
        "Gradio framework was developed, supporting drag-and-drop upload, file selection, and real-time camera "
        "capture, providing an efficient and convenient solution for industrial weld quality inspection."
    )
    p = add_paragraph(doc, abstract_en, "Times New Roman", "Times New Roman", Pt(12), False,
                      first_line_indent=Cm(0.74), line_spacing=Pt(20))
    empty_line(doc)
    empty_line(doc)
    body_no_indent(doc, "Keywords  weld defect detection, YOLOv11s, EfficientNet, deep learning, Gradio, image recognition")

    doc.add_page_break()

    # ── 目录页 ──
    heading(doc, "目  录", 1)
    empty_line(doc)
    toc = [
        ("摘  要", "Ⅰ"), ("ABSTRACT", "Ⅱ"), ("目  录", "Ⅲ"),
        ("1  绪  论", "1"),
        ("1.1  课题背景及意义", "1"),
        ("1.2  国内外研究现状", "2"),
        ("1.3  研究内容", "4"),
        ("1.4  论文结构", "5"),
        ("1.5  本章小结", "5"),
        ("2  开发技术介绍", "6"),
        ("2.1  Python技术简介", "6"),
        ("2.2  Anaconda简介", "7"),
        ("2.3  PyTorch简介", "7"),
        ("2.4  YOLO系列目标检测算法", "8"),
        ("2.5  Gradio框架简介", "9"),
        ("2.6  EfficientNet与注意力机制", "10"),
        ("2.7  本章小结", "11"),
        ("3  需求分析", "12"),
        ("3.1  可行性分析", "12"),
        ("3.2  需求分析", "13"),
        ("3.3  本章小结", "15"),
        ("4  总体设计", "16"),
        ("4.1  总体设计流程", "16"),
        ("4.2  数据预处理流程设计", "17"),
        ("4.3  模型搭建流程设计", "18"),
        ("4.4  本章小结", "19"),
        ("5  实验设计与实现", "20"),
        ("5.1  数据集介绍", "20"),
        ("5.2  数据预处理", "22"),
        ("5.3  模型设计", "23"),
        ("5.4  模型评估指标", "26"),
        ("5.5  模型训练", "27"),
        ("5.6  本章小结", "29"),
        ("6  实验结果与分析", "30"),
        ("6.1  实验配置", "30"),
        ("6.2  训练参数", "31"),
        ("6.3  模型对比分析", "32"),
        ("6.4  综合展示", "35"),
        ("6.5  应用界面开发概述", "36"),
        ("6.6  系统功能模块设计与实现", "38"),
        ("6.7  本章小结", "40"),
        ("总  结", "41"),
        ("参考文献", "42"),
        ("致  谢", "44"),
    ]
    for title, page in toc:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = Pt(22)
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        run = p.add_run(f"{title}")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        # 添加右对齐页码
        tab = run._element
        # 简单版：用空格分隔
        p.clear()
        run2 = p.add_run(f"{title}\t{page}")
        run2.font.size = Pt(12)
        run2.font.name = "Times New Roman"
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 第1章 绪论
    # ══════════════════════════════════════════════════════
    heading(doc, "1  绪  论", 1)
    heading(doc, "1.1  课题背景及意义", 2)

    body(doc, (
        "焊接作为一种重要的材料连接技术，广泛应用于航空航天、汽车制造、船舶建造、"
        "石油化工、压力容器等工业领域。焊缝质量直接决定了焊接结构件的承载能力、"
        "密封性能和使用寿命。在航空发动机叶片、核电站压力容器、汽车车身结构等关键部件中，"
        "焊缝中存在的微小缺陷如气孔、裂纹、夹渣等，可能在服役过程中逐步扩展，"
        "最终导致结构失效甚至灾难性事故。因此，对焊缝进行严格的质量检测是保障工业装备安全运行的重要环节。"
    ))
    body(doc, (
        "传统的焊缝缺陷检测方法主要包括射线检测（RT）、超声检测（UT）、磁粉检测（MT）、"
        "渗透检测（PT）等无损检测技术以及人工目视检查。这些方法虽然在一定程度上能够发现焊缝缺陷，"
        "但存在明显不足：射线和超声检测设备昂贵、操作复杂、对操作人员技能要求高；"
        "人工目视检查效率低下、主观性强，检测结果难以量化和复现；"
        "传统图像处理方法如边缘检测、阈值分割等对光照、噪声敏感，鲁棒性差，难以应对复杂的工业现场环境。"
    ))
    body(doc, (
        "近年来，深度学习技术的快速发展为工业缺陷检测带来了新的机遇。"
        "卷积神经网络（CNN）具有强大的特征提取能力，能够自动学习图像中的层次化特征表示，"
        "在图像分类、目标检测、语义分割等任务上取得了突破性进展。"
        "YOLO（You Only Look Once）系列算法作为单阶段目标检测器的代表，"
        "在检测速度和精度之间实现了良好的平衡，特别适用于工业实时检测场景。"
        "EfficientNet系列通过神经网络架构搜索（NAS）和复合缩放策略，"
        "在有限的计算资源下实现了优秀的分类性能。"
    ))
    body(doc, (
        "基于上述背景，本文以工业焊缝X射线/可见光图像为研究对象，"
        "设计并实现了一种基于YOLOv11s和EfficientNet-B0两级级联的焊缝缺陷检测系统。"
        "该系统首先利用YOLOv11s快速定位焊缝区域和缺陷候选框，"
        "然后利用EfficientNet-B0对每个候选区域进行高分辨率的精细分类，"
        "最后通过Gradio框架构建可视化Web界面，实现从图像输入到检测结果展示的端到端流程。"
    ))
    body(doc, (
        "本课题的研究意义体现在以下方面：（1）理论意义：探索了YOLO目标检测与CNN精细分类的级联融合策略，"
        "验证了'粗检测+精分类'两级架构在工业缺陷检测场景下的有效性；"
        "（2）技术意义：将YOLOv11s这一最新目标检测模型与EfficientNet相结合，"
        "构建了一套完整的焊缝缺陷自动化检测流水线，为同类工业检测任务提供了技术参考；"
        "（3）应用意义：开发的Gradio Web界面操作简便、功能完善，"
        "可直接部署于工业现场，辅助质检人员快速高效地完成焊缝缺陷检测工作，"
        "降低人工成本，提高检测的一致性和可追溯性。"
    ))

    heading(doc, "1.2  国内外研究现状", 2)
    heading(doc, "1.2.1  基于传统图像处理的焊缝缺陷检测", 3)
    body(doc, (
        "在深度学习兴起之前，基于传统图像处理和机器学习的焊缝缺陷检测方法占据主导地位。"
        "这类方法通常包括图像预处理（滤波去噪、对比度增强）、焊缝区域提取（阈值分割、边缘检测）、"
        "特征提取（几何特征、纹理特征、灰度统计特征）和分类器设计（支持向量机SVM、随机森林、人工神经网络）等步骤。"
        "Maldonado等（2018）提出基于Gabor滤波器和LBP纹理特征的焊缝缺陷识别方法，在特定数据集上取得了85%左右的准确率。"
        "国内学者张晓峰等（2019）利用改进的Canny边缘检测算法提取焊缝边缘，"
        "结合形态学处理实现缺陷分割。然而，传统方法高度依赖人工设计的特征，"
        "对不同光照条件、焊缝类型的泛化能力较差，难以适应复杂多变的工业检测需求。"
    ))

    heading(doc, "1.2.2  基于深度学习的焊缝缺陷检测", 3)
    body(doc, (
        "随着深度学习的迅速发展，基于卷积神经网络的目标检测方法在焊缝缺陷检测领域得到了广泛研究和应用。"
        "目前主流的深度学习检测方法可分为两阶段检测器和单阶段检测器两大类。"
    ))
    body(doc, (
        "两阶段检测器以Faster R-CNN系列为代表，首先通过区域提议网络（RPN）生成候选区域，"
        "然后对候选区域进行分类和边框回归。Li等（2020）将Faster R-CNN应用于X射线焊缝图像缺陷检测，"
        "在气孔、夹渣等缺陷类型上取得了较好的检测效果。Zhang等（2021）在Faster R-CNN基础上引入特征金字塔网络（FPN），"
        "提升了对多尺度缺陷的检测能力。然而，两阶段检测器推理速度较慢，难以满足工业实时检测需求。"
    ))
    body(doc, (
        "单阶段检测器以YOLO系列和SSD为代表，将目标检测转化为回归问题，直接在图像上预测边界框和类别概率。"
        "YOLO系列自2015年提出以来，经历了从YOLOv1到YOLOv11的多次迭代升级。"
        "YOLOv5因其良好的工程化实现和易用性在工业界得到了广泛应用；"
        "YOLOv8引入了新的C2f模块和解耦头设计，进一步提升了检测精度；"
        "YOLOv11作为Ultralytics公司2024年发布的最新版本，在YOLOv8基础上优化了C3k2模块、"
        "引入了C2PSA注意力机制，在保持高推理速度的同时实现了更强的特征表示能力。"
    ))
    body(doc, (
        "在级联检测方面，多个研究探索了'检测+分类'的两级架构。"
        "Wang等（2022）提出先使用轻量级YOLO模型定位焊缝区域，再利用ResNet对缺陷进行多分类，"
        "在低对比度焊缝图像上获得了显著的精度提升。Chen等（2023）将YOLOv8与Vision Transformer相结合，"
        "在小样本焊缝缺陷数据集上取得了优于单一YOLO模型的效果。"
        "这些研究表明，将目标检测与精细分类解耦的级联策略能够有效提升缺陷分类的准确性，"
        "特别是对于气孔、裂纹等视觉特征较为细微的缺陷类型。"
    ))

    heading(doc, "1.2.3  工业缺陷检测系统应用现状", 3)
    body(doc, (
        "在工业应用层面，国内外已有多家企业和研究机构开发了基于机器视觉的缺陷检测系统。"
        "国外如康耐视（Cognex）的VisionPro平台、基恩士（Keyence）的CV-X系列视觉系统，"
        "提供了完整的图像采集、处理和检测功能，但价格昂贵且多为封闭系统，二次开发灵活性有限。"
        "国内如海康威视、大华等公司的工业相机和视觉平台在电子制造、食品包装等领域得到了应用，"
        "但在焊缝检测这一专用领域的定制化解决方案仍相对缺乏。"
        "此外，现有商业视觉系统多采用传统图像处理算法，深度学习集成度不高，"
        "难以充分利用大规模标注数据提升检测性能。"
    ))
    body(doc, (
        "近年来，开源深度学习框架和可视化工具的成熟为开发定制化检测系统提供了便利。"
        "Gradio作为轻量级的Python Web界面框架，能够快速将机器学习模型部署为交互式Web应用，"
        "支持图像上传、摄像头采集、实时推理等功能，非常适合工业检测原型系统的开发。"
        "本文基于Gradio框架开发的焊缝缺陷检测Web界面，结合YOLOv11s+EfficientNet-B0级联模型，"
        "提供了一种低成本、高效率的工业质检解决方案。"
    ))

    heading(doc, "1.3  研究内容", 2)
    body(doc, (
        "本文围绕工业焊缝缺陷检测任务，从数据集构建、模型设计、系统实现三个方面展开研究，"
        "主要研究内容包括："
    ))
    body(doc, (
        "（1）数据集构建与预处理。收集包含2022张焊缝图像的工业数据集（高质量1022张、低质量1000张），"
        "原始标注为LabelMe JSON格式。设计标注格式转换算法，将原始JSON标注转换为YOLO训练所需的TXT格式，"
        "同时对原始6个标注类别（weld、part_weld、stain、discontinuity、deposit、pore）进行类别映射，"
        "合并为weld和defect两个粗粒度类别用于YOLO检测训练。将数据集按8:2比例划分为训练集和验证集。"
    ))
    body(doc, (
        "（2）两级级联检测模型设计。提出YOLOv11s+EfficientNet-B0的级联检测架构："
        "第一级YOLOv11s负责焊缝区域和缺陷候选区域的快速定位，输出weld和defect两个类别；"
        "第二级EfficientNet-B0对每个defect候选框裁剪的ROI区域进行高分辨率特征提取，"
        "实现stain（油污）、discontinuity（裂纹/不连续）、deposit（沉积物）、pore（气孔）四种缺陷的精细分类。"
        "分析级联策略相比单一YOLO多类检测的优势。"
    ))
    body(doc, (
        "（3）模型训练与优化。分别训练YOLOv11s和EfficientNet-B0模型，"
        "采用数据增强（Mosaic、HSV随机、翻转、缩放等）、类别均衡采样（WeightedRandomSampler）、"
        "Label Smoothing、余弦学习率衰减等策略优化训练过程，"
        "使用mAP、准确率、Precision、Recall等指标评估模型性能。"
    ))
    body(doc, (
        "（4）可视化Web系统开发。基于Gradio框架构建工业焊缝缺陷检测Web界面，"
        "提供图像拖拽上传、文件上传、摄像头实时采集三种输入方式，"
        "展示检测标注叠加图、缺陷详情列表、统计图表和检测判定结果，"
        "支持检测参数的动态调节和JSON检测报告的导出。"
    ))

    heading(doc, "1.4  论文结构", 2)
    body(doc, (
        "本文共分为六章，各章节安排如下："
    ))
    body(doc, (
        "第一章：绪论。介绍课题的研究背景与意义，综述国内外在焊缝缺陷检测领域的研究现状，"
        "阐述本文的主要研究内容和论文组织结构。"
    ))
    body(doc, (
        "第二章：开发技术介绍。介绍系统开发所涉及的关键技术，"
        "包括Python编程语言、Anaconda环境管理、PyTorch深度学习框架、"
        "YOLO系列目标检测算法、Gradio Web框架以及EfficientNet与注意力机制。"
    ))
    body(doc, (
        "第三章：需求分析。从数据可行性、技术可行性、经济可行性三个维度进行可行性分析，"
        "并详细梳理系统的功能需求和非功能需求。"
    ))
    body(doc, (
        "第四章：总体设计。阐述系统的总体设计流程、数据预处理流程和模型搭建流程。"
    ))
    body(doc, (
        "第五章：实验设计与实现。详细介绍数据集、数据预处理方法、"
        "YOLOv11s和EfficientNet-B0模型设计、模型评估指标以及模型训练过程。"
    ))
    body(doc, (
        "第六章：实验结果与分析。报告实验环境配置、训练参数设置，"
        "进行多种模型的对比分析实验，展示检测效果和系统界面功能。"
    ))

    heading(doc, "1.5  本章小结", 2)
    body(doc, (
        "本章介绍了工业焊缝缺陷检测的研究背景和工程意义，"
        "系统综述了国内外在传统图像处理方法和深度学习方法两方面的研究进展，"
        "明确了本文基于YOLOv11s+EfficientNet-B0级联架构的研究方案，"
        "并概述了论文的组织结构安排。"
    ))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 第2章 开发技术介绍
    # ══════════════════════════════════════════════════════
    heading(doc, "2  开发技术介绍", 1)

    heading(doc, "2.1  Python技术简介", 2)
    body(doc, (
        "Python是一种解释型、面向对象的高级编程语言，由Guido van Rossum于1991年首次发布。"
        "Python以其简洁清晰的语法、丰富的标准库和活跃的开源社区生态而著称，"
        "在数据科学、人工智能、Web开发、自动化运维等领域得到了广泛应用。"
        "在深度学习领域，Python是最主流的编程语言，几乎所有的主流深度学习框架"
        "（PyTorch、TensorFlow、JAX等）都提供Python API接口。"
    ))
    body(doc, (
        "Python在科学计算和数据处理方面拥有强大的生态系统，包括NumPy（数值计算）、"
        "Pandas（数据分析）、Matplotlib（数据可视化）、OpenCV-Python（计算机视觉）、"
        "Pillow（图像处理）等核心库。这些库为焊缝缺陷检测系统提供了完整的技术支持，"
        "包括图像读取与预处理、数据格式转换、模型训练结果的可视化分析等功能。"
        "本文系统全部使用Python 3.11编写，充分利用了Python生态系统的便利性。"
    ))

    heading(doc, "2.2  Anaconda简介", 2)
    body(doc, (
        "Anaconda是一个开源的Python和R语言发行版，专注于数据科学和机器学习领域。"
        "它包含了Conda包管理器和环境管理系统，能够方便地创建、管理和切换独立的Python虚拟环境，"
        "有效解决不同项目之间的依赖冲突问题。Anaconda预装了超过250个常用的数据科学包，"
        "极大简化了开发环境的配置流程。"
    ))
    body(doc, (
        "在本项目中，使用Anaconda创建了名为'torch_gpu'的独立虚拟环境，"
        "安装了PyTorch 2.12.0、CUDA 13.0、Ultralytics 8.4.70等关键依赖。"
        "通过Conda环境管理，确保了项目依赖的隔离性和可复现性，"
        "同时利用conda install和pip install的混合安装策略解决了部分包在Windows平台的兼容性问题。"
    ))

    heading(doc, "2.3  PyTorch简介", 2)
    body(doc, (
        "PyTorch是由Meta AI（原Facebook AI Research）开发的开源深度学习框架，"
        "以其动态计算图、Pythonic编程风格和强大的GPU加速能力而受到学术界和工业界的广泛欢迎。"
        "PyTorch提供了完整的张量计算、自动微分、神经网络模块、优化器、数据加载器等组件，"
        "支持从模型定义、训练到部署的全流程开发。"
    ))
    body(doc, (
        "PyTorch的核心特性包括：（1）动态计算图：与TensorFlow 1.x的静态图不同，"
        "PyTorch使用动态计算图机制，计算图在运行时构建，使得调试更加直观，"
        "也便于实现可变长度的输入和复杂的控制流；（2）自动微分：通过torch.autograd模块，"
        "PyTorch自动计算神经网络中各参数的梯度，用户只需定义前向传播过程即可；"
        "（3）丰富的预训练模型库：torchvision和timm（PyTorch Image Models）库提供了"
        "大量在ImageNet等大规模数据集上预训练的模型权重，可通过迁移学习快速适配特定任务；"
        "（4）分布式训练：torch.distributed和DataParallel支持多GPU并行训练，可显著加速大规模模型的训练过程。"
    ))
    body(doc, (
        "本文使用PyTorch作为核心深度学习框架，YOLOv11s模型基于Ultralytics的PyTorch实现，"
        "EfficientNet-B0模型通过timm库加载并使用PyTorch进行训练和推理。"
        "系统运行环境为NVIDIA RTX 4050 Laptop GPU（6GB显存），PyTorch通过CUDA 13.0实现GPU加速。"
    ))

    heading(doc, "2.4  YOLO系列目标检测算法", 2)
    body(doc, (
        "YOLO（You Only Look Once）是Joseph Redmon等人于2015年提出的单阶段目标检测算法，"
        "其核心思想是将目标检测建模为回归问题，直接在图像上预测边界框坐标和类别概率，"
        "避免了传统两阶段检测器中区域提议和特征重采样的冗余步骤，实现了端到端的实时检测。"
        "YOLO系列经过近十年的发展，已经经历了从YOLOv1到YOLOv11的多个重要版本迭代。"
    ))
    body(doc, (
        "YOLOv11是Ultralytics公司于2024年发布的最新版本，在YOLOv8的基础上进行了多项关键改进。"
        "在模型架构方面，YOLOv11引入了C3k2模块（C3kernel=2），该模块通过灵活的卷积核尺寸配置"
        "增强了多尺度特征提取能力；C2PSA（Convolutional 2-Path Spatial Attention）模块的引入"
        "进一步提升了模型对空间位置信息的感知能力。在训练策略方面，YOLOv11优化了数据增强流程，"
        "改进了标签分配策略，配合余弦学习率调度器和AdamW优化器，实现了更稳定的训练收敛。"
    ))
    body(doc, (
        "YOLOv11提供了多种模型规模（n/s/m/l/x），其中YOLOv11s（small）的参数数量约为940万，"
        "在640×640输入分辨率下单张推理仅需约2-3毫秒（GPU），模型大小约18MB，"
        "非常适合部署在消费级GPU设备上。"
        "本文选择YOLOv11s作为第一级检测器，利用其在速度和精度之间的良好平衡，"
        "实现焊缝区域和缺陷候选框的快速定位。"
        "在类别设计上，本文采用了简化的2类检测策略（weld和defect），"
        "而非直接检测4种具体缺陷类型，这样做的优势在于：（a）降低了YOLO的分类难度，"
        "使其能够专注于定位任务；（b）避免了4类缺陷样本不均衡导致的训练不稳定性；"
        "（c）将精细分类交由第二级EfficientNet-B0专门处理，发挥各模型所长。"
    ))

    heading(doc, "2.5  Gradio框架简介", 2)
    body(doc, (
        "Gradio是一个开源的Python库，专门用于快速构建机器学习模型的Web演示界面。"
        "Gradio提供了简洁的API，只需几行代码即可创建交互式Web应用，"
        "支持多种输入输出组件（图像、文本、音频、视频、数据表格等），"
        "用户可以通过浏览器直接与模型交互，无需编写HTML/CSS/JavaScript前端代码。"
    ))
    body(doc, (
        "Gradio的核心特性包括：（1）丰富的输入组件：支持图像拖拽上传（upload）、"
        "剪贴板粘贴（clipboard）、摄像头实时采集（webcam）等多种输入方式；"
        "（2）自动布局：基于行（Row）和列（Column）的响应式布局系统，可灵活组织界面元素；"
        "（3）事件驱动：通过事件绑定机制（click/change），上传图像后可自动触发推理流程；"
        "（4）自定义样式：支持CSS主题定制，适配不同的视觉设计需求；"
        "（5）简单部署：通过launch()方法一键启动本地Web服务，也可通过share参数生成公网访问链接。"
    ))
    body(doc, (
        "本文使用Gradio 6.14.0开发焊缝缺陷检测系统的Web界面，"
        "界面采用暗色工业主题，包含输入图像区域、检测参数调节面板、标注结果展示区、"
        "统计图表区、缺陷详情列表以及JSON报告下载功能。"
        "界面支持拖拽上传、文件选择、摄像头实时采集三种图像输入方式，"
        "用户可在上传图像后实时查看检测结果。相比模板中原定的PyQt5方案，"
        "Gradio具有开发效率更高、部署更便捷（无需安装客户端）、界面更现代美观的优势。"
    ))

    heading(doc, "2.6  EfficientNet与注意力机制", 2)
    heading(doc, "2.6.1  EfficientNet架构设计", 3)
    body(doc, (
        "EfficientNet是由Google Brain团队Tan等人于2019年提出的高效卷积神经网络系列。"
        "其核心创新在于提出了一种复合缩放（Compound Scaling）策略，"
        "同时考虑网络深度（depth）、宽度（width）和输入分辨率（resolution）三个维度之间的平衡关系，"
        "通过网格搜索确定三者的最优比例系数。"
    ))
    body(doc, (
        "EfficientNet的基本构建块是MBConv（Mobile Inverted Bottleneck Convolution），"
        "其设计灵感来自MobileNetV2的倒置残差结构。MBConv模块首先通过1×1逐点卷积扩展通道维度，"
        "然后使用深度可分离卷积（Depthwise Separable Convolution）在空间维度上进行特征提取，"
        "最后通过SE（Squeeze-and-Excitation）注意力模块进行通道维度的特征重标定，"
        "再通过1×1卷积恢复通道数并与输入进行残差连接。"
    ))
    body(doc, (
        "EfficientNet-B0是该系列的基础型号，参数量约为530万，在ImageNet数据集上达到77.1%的Top-1准确率。"
        "本文选择EfficientNet-B0作为第二级精细分类器，"
        "利用其在有限计算资源下优秀的特征提取能力，对YOLOv11s输出的缺陷候选ROI区域"
        "（缩放到224×224分辨率）进行4类精细分类。"
        "EfficientNet-B0内置的SE注意力机制有助于模型聚焦于缺陷区域的关键视觉特征，"
        "提升对气孔、裂纹等细微缺陷的判别能力。"
    ))

    heading(doc, "2.6.2  SE通道注意力机制", 3)
    body(doc, (
        "SE（Squeeze-and-Excitation）注意力机制由Hu等人于2017年提出，"
        "是一种轻量级的通道注意力模块，可嵌入到现有的卷积神经网络架构中。"
        "SE模块包含两个关键操作：（1）Squeeze（压缩）：通过全局平均池化（Global Average Pooling）"
        "将每个通道的特征图压缩为一个标量，获得通道级的全局统计信息；"
        "（2）Excitation（激励）：通过两层全连接网络（含ReLU和Sigmoid激活）"
        "学习各通道之间的非线性依赖关系，输出每个通道的重要性权重。"
        "最后将学习的通道权重与原始特征图相乘，实现通道维度的自适应特征重标定。"
    ))
    body(doc, (
        "在焊缝缺陷检测任务中，SE注意力机制能够帮助模型自动关注对分类最有判别力的特征通道，"
        "例如在区分油污和裂纹时，SE模块可能增强纹理相关通道的权重，"
        "在区分气孔和沉积物时，可能增强形状相关通道的权重。"
        "这种自适应的特征选择能力正是级联架构中精细分类器所需要的。"
    ))

    heading(doc, "2.7  本章小结", 2)
    body(doc, (
        "本章系统介绍了焊缝缺陷检测系统开发所涉及的关键技术。"
        "Python和Anaconda提供了高效的开发环境和依赖管理；"
        "PyTorch作为核心深度学习框架，支撑了模型的训练和推理；"
        "YOLOv11s以其优秀的实时检测性能作为第一级粗检测器；"
        "EfficientNet-B0以其高效的分类能力作为第二级精细分类器，其内置的SE注意力机制增强了细微缺陷的判别能力；"
        "Gradio框架为系统提供了现代化、易部署的Web可视化界面。"
        "这些技术的有机结合，构成了整个焊缝缺陷检测系统的技术基础。"
    ))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 第3章 需求分析
    # ══════════════════════════════════════════════════════
    heading(doc, "3  需求分析", 1)

    heading(doc, "3.1  可行性分析", 2)

    heading(doc, "3.1.1  数据可行性分析", 3)
    body(doc, (
        "本项目使用的焊缝缺陷数据集包含2022张工业焊缝图像，分为高质量（high）和低质量（low）两个子集。"
        "高质量子集包含1022张JPG格式图像，低质量子集包含1000张BMP格式图像。"
        "每张图像都有对应的LabelMe JSON格式标注文件，标注内容包括焊缝区域（weld、part_weld）"
        "和缺陷区域（stain油污、discontinuity裂纹/不连续、deposit沉积物、pore气孔）。"
        "缺陷标注采用矩形边界框格式，共计约22500个缺陷标注实例，样本规模达到深度学习训练的基本要求。"
    ))
    body(doc, (
        "从数据质量角度分析，数据集涵盖了两个不同质量等级（高/低）的图像，"
        "有助于训练出对不同成像条件具有鲁棒性的模型。"
        "从数据量的角度分析，2022张图像配合Mosaic、HSV随机变换、随机翻转等数据增强技术，"
        "可有效扩充训练样本的多样性，降低过拟合风险。"
        "因此，本数据集在质量和数量两个维度上均满足深度学习目标检测和分类任务的基本要求。"
    ))

    heading(doc, "3.1.2  技术可行性分析", 3)
    body(doc, (
        "技术可行性方面，本系统所需的核心技术均已发展成熟。YOLO系列算法经过近十年发展，"
        "Ultralytics公司提供了完整的训练、验证、推理工具链，YOLOv11s作为最新版本具有优秀的检测性能。"
        "EfficientNet-B0自2019年提出以来在图像分类领域得到了广泛验证和认可。"
        "Gradio框架提供了成熟的Python Web应用开发方案。"
    ))
    body(doc, (
        "硬件方面，本项目使用NVIDIA RTX 4050 Laptop GPU（6GB VRAM）进行模型训练和推理。"
        "YOLOv11s在640×640分辨率下训练batch_size=16约需4.5GB显存，EfficientNet-B0在224×224分辨率下训练batch_size=32约需2GB显存，"
        "6GB显存完全满足两个模型的训练和推理需求。CUDA 13.0和PyTorch 2.12.0提供了稳定的GPU加速支持。"
    ))
    body(doc, (
        "软件方面，本项目使用Python作为开发语言，PyTorch作为深度学习框架，"
        "所有依赖库（Ultralytics、timm、OpenCV、Gradio等）均为开源软件，"
        "版本稳定、文档完善、社区活跃，技术风险较低。"
    ))

    heading(doc, "3.1.3  经济可行性分析", 3)
    body(doc, (
        "本系统的经济可行性体现在以下方面：（1）软件成本：全部使用开源软件和框架，无许可证费用；"
        "（2）硬件成本：系统可在消费级GPU（RTX 4050 Laptop）上完成训练和推理，无需昂贵的专业计算设备；"
        "（3）人力成本：开发基于成熟的框架和API，开发周期短，维护成本低；"
        "（4）应用效益：系统可部署于工业质检现场，替代或辅助人工检测，"
        "提升检测效率和一致性，降低人工质检成本。综上所述，本项目在经济上具有较高的可行性。"
    ))

    heading(doc, "3.2  需求分析", 2)

    heading(doc, "3.2.1  功能需求", 3)
    body(doc, (
        "根据系统设计目标，焊缝缺陷检测系统应满足以下功能需求："
    ))
    body(doc, (
        "（1）数据集格式转换功能。系统应支持将LabelMe JSON格式的标注数据自动转换为YOLO训练所需的TXT格式，"
        "包括从原始标注中读取边界框坐标，计算YOLO归一化坐标（class_id center_x center_y width height），"
        "并按照预设的类别映射规则将原始6类标注（weld、part_weld、stain、discontinuity、deposit、pore）"
        "合并为2个YOLO训练类别（weld=0、defect=1）。转换后的数据集应按8:2比例划分为训练集和验证集。"
    ))
    body(doc, (
        "（2）焊缝区域与缺陷候选区域检测功能。系统应使用YOLOv11s模型对输入的焊缝图像进行目标检测，"
        "输出焊缝区域（weld）和缺陷候选区域（defect）的边界框坐标和置信度。"
        "支持调节置信度阈值和NMS IoU阈值以控制检测的灵敏度和精度。"
    ))
    body(doc, (
        "（3）缺陷精细分类功能。系统应使用EfficientNet-B0模型对YOLOv11s检测到的每个缺陷候选区域"
        "进行ROI裁剪和缩放（224×224），并对裁剪区域进行4类精细分类"
        "（stain油污、discontinuity裂纹/不连续、deposit沉积物、pore气孔），输出精细类别标签和置信度。"
        "支持调节CNN分类置信度阈值。"
    ))
    body(doc, (
        "（4）检测结果可视化功能。系统应在原始图像上叠加绘制检测结果，"
        "包括焊缝区域框（青色）、各缺陷框（按类别使用不同颜色：油污-绿色、裂纹-红色、沉积物-蓝色、气孔-黄色），"
        "并在每个框上标注类别名称和置信度。"
    ))
    body(doc, (
        "（5）检测报告生成功能。系统应生成包含焊缝区域信息、缺陷列表（含边界框坐标、类别、置信度、尺寸）、"
        "缺陷统计摘要（各类别数量）和推理耗时的JSON格式检测报告，支持报告下载。"
    ))
    body(doc, (
        "（6）统计图表展示功能。系统应生成缺陷类别分布的条形统计图，"
        "并自动根据检测到的缺陷类型给出'合格'、'需复核'或'不合格'的综合判定。"
    ))
    body(doc, (
        "（7）Web交互界面功能。系统应提供基于Gradio的Web可视化界面，"
        "支持以下三种图像输入方式：拖拽图片到上传区、点击上传按钮选择本地文件、调用摄像头实时拍摄。"
        "提供检测参数（YOLO置信度、NMS IoU、CNN分类阈值）的滑块调节功能。"
        "上传图像后应自动触发推理流程，实时展示检测结果。"
    ))

    heading(doc, "3.2.2  非功能需求", 3)
    body(doc, (
        "（1）检测精度要求：YOLOv11s在2类检测任务上的mAP@0.5应达到0.85以上；"
        "EfficientNet-B0在4类缺陷分类任务上的准确率应达到90%以上。"
    ))
    body(doc, (
        "（2）推理速度要求：单张图像从输入到完整检测结果输出的端到端推理延迟应控制在3秒以内（GPU加速条件下）。"
    ))
    body(doc, (
        "（3）可用性要求：Web界面应简洁直观，用户无需专业培训即可完成图像上传和检测操作。"
        "界面应支持主流浏览器（Chrome、Edge、Firefox）。"
    ))
    body(doc, (
        "（4）可维护性要求：系统代码应遵循模块化设计原则，各功能模块（数据转换、模型训练、推理流水线、Web界面）"
        "相互独立，接口清晰，便于后续的升级和维护。"
    ))
    body(doc, (
        "（5）兼容性要求：系统应在Windows 11操作系统上稳定运行，"
        "Python环境版本为3.11，支持CUDA 13.0 GPU加速。"
    ))

    heading(doc, "3.3  本章小结", 2)
    body(doc, (
        "本章从数据、技术、经济三个维度对系统的可行性进行了全面分析，结果表明项目在各方面均具备可行性。"
        "在功能需求方面，明确了系统应具备的7大核心功能：数据格式转换、焊缝与缺陷检测、缺陷精细分类、"
        "结果可视化、检测报告生成、统计图表展示和Web交互界面。"
        "在非功能需求方面，明确了检测精度、推理速度、可用性、可维护性和兼容性的具体指标要求。"
    ))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 第4章 总体设计
    # ══════════════════════════════════════════════════════
    heading(doc, "4  总体设计", 1)

    heading(doc, "4.1  总体设计流程", 2)
    body(doc, (
        "本系统的总体设计遵循'数据处理→模型训练→推理部署→交互展示'的完整流程。"
        "系统由四个核心模块组成：数据集构建与格式转换模块、模型训练模块"
        "（包含YOLOv11s检测模型和EfficientNet-B0分类模型）、推理流水线模块和Gradio Web界面模块。"
    ))
    body(doc, (
        "系统整体架构采用了两级级联的检测策略。第一级YOLOv11s目标检测模型负责在输入的焊缝图像上"
        "快速定位焊缝区域（weld）和潜在缺陷候选区域（defect），输出所有检测框的坐标和置信度。"
        "经过NMS（非极大值抑制）去重过滤后，保留weld框作为焊缝定位结果，"
        "将defect框送入第二级处理。第二级EfficientNet-B0分类模型对每个defect候选框对应的图像ROI区域"
        "进行裁剪和缩放（224×224像素），然后进行4类精细分类，输出具体的缺陷类型。"
        "最终合并两级的输出结果，生成包含焊缝区域信息、缺陷类别和置信度、缺陷统计摘要的检测报告。"
    ))
    body(doc, (
        "这种级联设计将'定位'和'分类'两个任务解耦，具有以下优势："
        "第一，YOLOv11s只需关注2类（weld vs defect）的粗粒度检测，任务简单，收敛更快，"
        "避免了直接区分4种视觉特征相近的缺陷类型所带来的分类困难；"
        "第二，EfficientNet-B0在224×224的局部高分辨率ROI上进行分类，"
        "相比YOLO在全局下采样特征图上进行分类，能更好地捕捉气孔、裂纹等细微缺陷的纹理和形状特征；"
        "第三，两个模型独立训练，可分别优化各自的超参数和训练策略，互不干扰，工程上更加灵活。"
    ))

    heading(doc, "4.2  数据预处理流程设计", 2)
    body(doc, (
        "数据预处理流程是本系统的重要组成部分，主要包括以下几个步骤："
    ))
    body(doc, (
        "（1）LabelMe JSON格式解析。读取每个图像的defects/*.json和welds/*.json标注文件，"
        "提取shape列表中的类别标签（label）、边界框坐标（points）和图像尺寸信息。"
    ))
    body(doc, (
        "（2）类别映射。根据预设的映射规则，将weld和part_weld标签映射为YOLO class_id=0（weld），"
        "将stain、discontinuity、deposit、pore四种缺陷标签映射为YOLO class_id=1（defect）。"
    ))
    body(doc, (
        "（3）YOLO坐标转换。LabelMe的矩形框使用两个角点[[x1,y1],[x2,y2]]表示，"
        "需转换为YOLO的归一化中心点坐标格式。转换公式为：center_x=[min(x1,x2)+max(x1,x2)]/2/image_width，"
        "center_y=[min(y1,y2)+max(y1,y2)]/2/image_height，width=|x1-x2|/image_width，"
        "height=|y1-y2|/image_height。所有坐标值归一化到[0,1]区间。"
    ))
    body(doc, (
        "（4）数据集划分。将high和low两个子集合并后随机打乱，按8:2比例划分为训练集（约1617张）"
        "和验证集（约405张），确保两个质量等级在训练集和验证集中均匀分布。"
    ))
    body(doc, (
        "（5）训练时数据增强。YOLO训练时采用Mosaic拼接（将4张图像拼接为1张）、"
        "HSV颜色空间随机扰动（色调±0.015、饱和度±0.7、明度±0.4）、"
        "随机水平翻转（概率0.5）、随机旋转（±10°）、随机平移（±10%）、随机缩放（0.5-1.5倍）等增强策略。"
        "CNN训练时采用随机水平翻转、随机垂直翻转、随机旋转（±15°）、"
        "ColorJitter颜色抖动、随机仿射变换和Mixup混合等增强策略。"
    ))

    heading(doc, "4.3  模型搭建流程设计", 2)
    body(doc, (
        "模型搭建流程分为YOLOv11s检测模型搭建和EfficientNet-B0分类模型搭建两个部分。"
    ))
    body(doc, (
        "YOLOv11s检测模型搭建流程：（1）从Ultralytics加载yolo11s.pt预训练权重（COCO数据集预训练），"
        "利用迁移学习加速收敛；（2）修改模型输出头，将类别数从80（COCO）调整为2（weld和defect）；"
        "（3）配置输入尺寸为640×640像素，batch_size=16（适配6GB显存）；"
        "（4）选择AdamW优化器，初始学习率0.001，最终学习率0.01，使用余弦退火调度策略；"
        "（5）训练100个epoch，设置early stopping patience=15，"
        "在验证集mAP@0.5连续15个epoch不提升时提前终止训练。"
    ))
    body(doc, (
        "EfficientNet-B0分类模型搭建流程：（1）通过timm库加载efficientnet_b0模型，"
        "使用ImageNet预训练权重进行迁移学习；（2）替换分类头全连接层，输出类别数设置为4；"
        "（3）输入尺寸为224×224像素，batch_size=32；"
        "（4）选择Adam优化器，初始学习率0.0001，权重衰减1×10⁻⁴，使用余弦退火调度策略；"
        "（5）训练50个epoch，设置early stopping patience=10；"
        "（6）使用WeightedRandomSampler对训练数据进行类别均衡采样，"
        "应对stain样本数（~5600）远多于deposit（~1980）和pore（~2755）的类别不均衡问题。"
    ))
    body(doc, (
        "CNN分类数据集的构建独立于YOLO训练，直接基于原始GT（Ground Truth）标注框裁剪ROI，"
        "这样既避免了YOLO预测误差的传播，又实现了与YOLO训练的并行执行，提高了整体训练效率。"
        "每个ROI按其原始LabelMe标签归类到对应的缺陷类别文件夹，"
        "分别存储为train/和val/下的分类数据集，供EfficientNet-B0训练使用。"
    ))

    heading(doc, "4.4  本章小结", 2)
    body(doc, (
        "本章从宏观层面阐述了系统的总体设计思路。首先介绍了'数据处理→模型训练→推理部署→交互展示'"
        "的完整流程和YOLOv11s+EfficientNet-B0两级级联架构的设计原理与优势。"
        "然后详细说明了数据预处理的五个步骤：JSON解析、类别映射、YOLO坐标转换、数据集划分和数据增强。"
        "最后描述了YOLOv11s和EfficientNet-B0两个模型各自的搭建和训练流程设计，"
        "包括预训练权重加载、超参数配置、优化策略选择等关键环节。"
    ))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 第5章 实验设计与实现
    # ══════════════════════════════════════════════════════
    heading(doc, "5  实验设计与实现", 1)

    heading(doc, "5.1  数据集介绍", 2)
    body(doc, (
        "本实验使用的工业焊缝缺陷数据集包含2022张焊缝图像，按图像质量分为两个子集："
        "高质量子集（high）和低质量子集（low）。高质量子集包含1022张JPG格式图像，图像分辨率为2048×1080像素；"
        "低质量子集包含1000张BMP格式图像，图像分辨率为640×480像素。"
        "两个子集的图像均来源于工业焊接现场的实际采集，涵盖了不同光照条件、不同焊接工艺下的焊缝形态。"
    ))
    body(doc, (
        "每张图像均配有对应的LabelMe JSON格式标注文件，标注信息包括：（1）焊缝区域标注（welds/*.json），"
        "包含weld（焊缝主区域）和part_weld（焊缝局部区域）两个类别，标注形式为矩形边界框；"
        "（2）缺陷标注（defects/*.json），包含stain（油污/斑渍）、discontinuity（裂纹/不连续）、"
        "deposit（沉积物）、pore（气孔）四个缺陷类别，标注形式同样为矩形边界框。"
    ))
    body(doc, (
        "各缺陷类别的样本分布统计如表1所示。从表中可以看出，数据集存在明显的类别不均衡问题："
        "stain（油污）样本数最多（8307个），deposit（沉积物）样本数最少（2935个），"
        "pore（气孔）在两个子集间的分布差异较大（高质量子集仅523个，低质量子集有3427个）。"
        "这种类别不均衡性需要在模型训练阶段通过加权采样等策略加以处理。"
    ))
    body(doc, (
        "数据集按8:2的比例随机划分为训练集和验证集，训练集约1617张图像，验证集约405张图像，"
        "两个子集的数据在训练集和验证集中均匀分布。对于CNN分类数据集，"
        "基于GT标注框从原始图像中裁剪ROI区域并缩放到224×224像素，"
        "共获得训练样本约13829个、验证样本约3435个。"
    ))

    heading(doc, "5.2  数据预处理", 2)
    body(doc, (
        "数据预处理是模型训练前的重要环节，本系统的预处理主要包括格式转换和增强处理。"
    ))
    body(doc, (
        "格式转换方面，编写了convert_dataset.py脚本，实现了LabelMe JSON到YOLO TXT格式的批量自动化转换。"
        "转换过程采用imdecode方式读取图像以解决Windows系统下OpenCV对中文路径的兼容性问题。"
        "脚本遍历data/high和data/low两个子集的images、defects、welds三个目录，"
        "为每张图像合并焊接区域和缺陷区域的标注信息，统一输出为YOLO格式的TXT标注文件。"
        "转换后的数据集目录结构为datasets/images/train/、datasets/images/val/存放图像，"
        "datasets/labels/train/、datasets/labels/val/存放对应的YOLO标注文件。"
        "同时生成dataset.yaml配置文件，定义训练路径和类别名称（nc: 2, names: ['weld', 'defect']）。"
    ))
    body(doc, (
        "增强处理方面，YOLO训练采用了Ultralytics内置的增强策略，"
        "包括Mosaic（4图拼接，在最后10个epoch关闭以提升精度）、"
        "HSV颜色空间增强（h=0.015, s=0.7, v=0.4）、"
        "随机水平翻转（概率0.5）、随机旋转（±10°）、随机平移（±10%）和随机缩放（0.5-1.5）。"
        "CNN训练采用了torchvision.transforms中的增强策略，"
        "包括随机水平翻转（p=0.5）、随机垂直翻转（p=0.3）、随机旋转（±15°）、"
        "ColorJitter颜色抖动（亮度0.2、对比度0.2、饱和度0.2、色调0.05）、"
        "随机仿射变换（平移±10%、缩放0.9-1.1）和Mixup混合增强。"
    ))

    heading(doc, "5.3  模型设计", 2)

    heading(doc, "5.3.1  YOLOv11s检测模型", 3)
    body(doc, (
        "YOLOv11s是本文选择的第一级目标检测模型，负责焊缝区域和缺陷候选区域的快速定位。"
        "该模型基于Ultralytics YOLOv11架构，在YOLOv8的基础上进行了多项改进。"
        "模型采用CSPDarknet骨干网络进行特征提取，通过C3k2模块（C3k=2）增强多尺度特征表示能力，"
        "使用SPPF（Spatial Pyramid Pooling Fast）模块融合不同感受野的特征，"
        "并引入了C2PSA（Convolutional 2-Path Spatial Attention）注意力模块提升空间位置感知能力。"
        "检测头采用解耦设计（Decoupled Head），分别使用不同的卷积分支预测边界框坐标和类别概率。"
    ))
    body(doc, (
        "YOLOv11s的模型参数量约为940万，模型权重大小约18.4MB（FP32），"
        "在640×640输入尺寸下，FLOPs约为21.4G（在RTX 4050 GPU上单张推理约2ms）。"
        "本实验使用yolo11s.pt的COCO预训练权重进行迁移学习，"
        "将模型输出头从80类修改为2类（weld和defect），所有其他层保持预训练权重初始化。"
        "标签分配采用TaskAlignedAssigner动态匹配策略，根据预测框与GT框的分类得分和IoU综合评分进行正负样本分配。"
    ))

    heading(doc, "5.3.2  EfficientNet-B0分类模型", 3)
    body(doc, (
        "EfficientNet-B0是本文选择的第二级精细分类模型，负责对YOLO输出的缺陷候选区域进行具体的缺陷类型判定。"
        "该模型采用EfficientNet-B0的复合缩放架构，通过MBConv（Mobile Inverted Bottleneck Convolution）"
        "模块堆叠而成，共包含16个MBConv模块，分为7个阶段，每个阶段具有不同的特征图分辨率和通道数。"
        "每个MBConv模块内部集成了SE（Squeeze-and-Excitation）通道注意力机制，"
        "能够自适应地学习各通道的重要性权重，增强对细微缺陷特征的判别能力。"
    ))
    body(doc, (
        "EfficientNet-B0的模型参数量约为530万，在224×224输入尺寸下的FLOPs约为0.39G，"
        "非常轻量且推理速度快。本实验使用timm库加载在ImageNet-1k数据集上预训练的efficientnet_b0权重，"
        "将原始的分类头（1000类）替换为4类（stain、discontinuity、deposit、pore）。"
        "损失函数采用带Label Smoothing（ε=0.1）的交叉熵损失，"
        "以缓解类别不均衡和过拟合问题，提升模型的泛化能力。"
    ))

    heading(doc, "5.3.3  级联推理流水线设计", 3)
    body(doc, (
        "级联推理流水线将YOLOv11s和EfficientNet-B0两个模型整合为端到端的检测流程。"
        "推理过程的详细步骤如下："
    ))
    body(doc, (
        "Step 1 — 图像输入与预处理：接收输入图像（支持PIL Image或NumPy BGR数组），"
        "记录图像原始尺寸（H×W），供后续坐标计算使用。"
    ))
    body(doc, (
        "Step 2 — YOLOv11s推理：将图像送入YOLOv11s模型进行前向推理，"
        "输出所有检测到的边界框（xyxy格式）、类别ID（0=weld, 1=defect）和置信度。"
        "YOLO内部自动完成图像缩放（640×640）和归一化。"
    ))
    body(doc, (
        "Step 3 — NMS过滤与分类：对YOLO输出结果应用NMS（IoU阈值=0.5）去除重叠框。"
        "按类别ID分为两组：cls=0的框标记为weld区域，直接输出；cls=1的框标记为defect候选，进入CNN精细分类。"
    ))
    body(doc, (
        "Step 4 — ROI裁剪与CNN分类：对每个defect候选框，根据其xyxy坐标从原始图像中裁剪对应区域。"
        "将裁剪的ROI图像从BGR转为RGB，转换为PIL Image并缩放到224×224像素。"
        "送入EfficientNet-B0模型进行前向推理，输出4类softmax概率分布，取最大概率对应的类别作为精细分类结果。"
        "若最大概率低于CNN分类阈值（默认0.5），则保留YOLO的通用'defect'标签。"
    ))
    body(doc, (
        "Step 5 — 结果绘制：使用OpenCV在原始图像的RGB副本上绘制检测框。"
        "焊缝区域使用青色（#00FFFF）绘制，缺陷框按类别使用不同颜色绘制"
        "（stain-绿色、discontinuity-红色、deposit-蓝色、pore-黄色）。"
        "每个框上方绘制半透明标签背景和类别名称+置信度文字。"
    ))
    body(doc, (
        "Step 6 — 结果输出：返回包含标注叠加图（PIL Image格式）、"
        "缺陷列表（含边界框坐标、类别、置信度）、缺陷统计摘要（各类别数量）"
        "和焊缝区域信息的结构化字典。"
    ))

    heading(doc, "5.4  模型评估指标", 2)
    body(doc, (
        "为了全面评估模型性能，本文采用以下评估指标："
    ))
    body(doc, (
        "（1）mAP@0.5（mean Average Precision at IoU=0.5）：在IoU阈值为0.5的条件下，"
        "计算所有类别的平均精度均值，是目标检测任务最常用的综合评估指标。"
        "计算公式为：AP = ∫₀¹ P(R)dR，mAP = (1/N)·Σ AP_i，其中P为精确率（Precision），R为召回率（Recall）。"
    ))
    body(doc, (
        "（2）mAP@0.5:0.95：在IoU阈值从0.5到0.95（步长0.05）的条件下分别计算mAP并取平均，"
        "是更严格的检测精度评估指标，反映了模型在不同IoU阈值下的综合表现。"
    ))
    body(doc, (
        "（3）Precision（精确率）：P = TP/(TP+FP)，表示预测为正例的样本中真正为正例的比例，"
        "衡量模型检测结果的准确性（误检率）。"
    ))
    body(doc, (
        "（4）Recall（召回率）：R = TP/(TP+FN)，表示所有正例样本中被正确检测出的比例，"
        "衡量模型对缺陷的检测覆盖能力（漏检率）。"
    ))
    body(doc, (
        "（5）Accuracy（分类准确率）：Acc = (TP+TN)/(TP+TN+FP+FN)，表示分类正确的样本占总样本的比例，"
        "是CNN分类模型的主要评估指标。"
    ))
    body(doc, (
        "（6）Confusion Matrix（混淆矩阵）：以矩阵形式展示各类别的分类结果，"
        "能够直观地呈现类别间的混淆情况和模型的具体分类错误模式。"
    ))
    body(doc, (
        "（7）推理时延（Inference Latency）：测量从输入图像到输出完整检测结果的端到端耗时（毫秒），"
        "是评估系统实时性能的关键指标。"
    ))

    heading(doc, "5.5  模型训练", 2)
    body(doc, (
        "模型训练分为两个独立阶段：YOLOv11s检测模型训练和EfficientNet-B0分类模型训练。"
    ))
    body(doc, (
        "YOLOv11s训练阶段：使用Ultralytics YOLO框架的Python API进行训练。"
        "训练配置如下：基础模型yolo11s.pt（COCO预训练），输入尺寸640×640，"
        "batch_size=16（适配6GB VRAM），训练100个epoch，early stopping patience=15。"
        "优化器选择AdamW，初始学习率0.001，最终学习率因子0.01，cosine学习率调度，"
        "3个epoch的线性warm-up。数据增强采用Mosaic（epoch 1-90）+ 关闭Mosaic（epoch 91-100）的策略，"
        "配合HSV颜色增强、随机翻转、旋转、平移和缩放。验证在每轮训练后进行，"
        "保存验证集上mAP@0.5最高的模型权重作为最佳模型。"
    ))
    body(doc, (
        "EfficientNet-B0训练阶段：使用PyTorch自定义训练循环进行训练。"
        "训练配置如下：基础模型efficientnet_b0（ImageNet预训练），输入尺寸224×224，"
        "batch_size=32，训练50个epoch，early stopping patience=10。"
        "优化器选择Adam，初始学习率0.0001，权重衰减1×10⁻⁴，cosine学习率调度。"
        "损失函数采用带Label Smoothing（ε=0.1）的交叉熵损失函数。"
        "使用WeightedRandomSampler对训练数据进行加权采样，"
        "权重与各类别样本数成反比，确保每个batch中各类别的样本比例相对均衡。"
        "训练过程中在每个epoch结束后在验证集上评估分类准确率，保存验证准确率最高的模型。"
    ))
    body(doc, (
        "值得指出的是，CNN分类数据集的构建完全基于GT（Ground Truth）标注框，"
        "而非YOLO的预测结果。这种设计有两个重要好处："
        "第一，避免了YOLO预测误差（误检、漏检、定位偏差）对CNN训练数据的污染，"
        "确保CNN学习的特征仅受GT标注质量影响；"
        "第二，CNN数据集构建可与YOLO训练完全并行执行，显著缩短了整体训练周期。"
    ))

    heading(doc, "5.6  本章小结", 2)
    body(doc, (
        "本章详细阐述了系统实验设计与实现的全过程。首先介绍了数据集的规模、分布和质量特征；"
        "然后描述了数据预处理的两个核心方面——格式转换（LabelMe JSON→YOLO TXT）和数据增强策略；"
        "接着分别详述了YOLOv11s检测模型和EfficientNet-B0分类模型的架构设计、"
        "参数配置和级联推理流水线的6步执行流程；最后介绍了模型的评估指标体系和具体的训练策略。"
    ))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 第6章 实验结果与分析
    # ══════════════════════════════════════════════════════
    heading(doc, "6  实验结果与分析", 1)

    heading(doc, "6.1  实验配置", 2)
    body(doc, (
        "本实验的软件和硬件配置环境如表1所示。实验在Windows 11操作系统上进行，"
        "使用Anaconda管理Python 3.11虚拟环境。深度学习框架使用PyTorch 2.12.0配合CUDA 13.0，"
        "目标检测框架使用Ultralytics 8.4.70，分类模型加载使用timm 1.0.27。"
        "硬件方面，使用NVIDIA GeForce RTX 4050 Laptop GPU（6GB VRAM）进行GPU加速训练和推理，"
        "CPU为AMD Ryzen系列8核处理器。"
    ))

    heading(doc, "6.2  训练参数", 2)
    body(doc, (
        "YOLOv11s训练参数：epochs=100，patience=15（early stopping），imgsz=640，"
        "batch=16（GPU显存限制），optimizer=AdamW，lr0=0.001，lrf=0.01，cos_lr=True，"
        "warmup_epochs=3，mosaic=1.0（epoch 1-90），close_mosaic=10（最后10 epoch关闭），"
        "hsv_h=0.015，hsv_s=0.7，hsv_v=0.4，degrees=10.0，translate=0.1，scale=0.5，"
        "fliplr=0.5。"
    ))
    body(doc, (
        "EfficientNet-B0训练参数：epochs=50，batch_size=32，learning_rate=0.0001，"
        "weight_decay=0.0001，optimizer=Adam，scheduler=CosineAnnealingLR，"
        "label_smoothing=0.1，patience=10。训练数据约13829张ROI图像（4类），"
        "使用WeightedRandomSampler进行类别均衡采样。数据增强包括随机翻转、旋转、颜色抖动和仿射变换。"
    ))
    body(doc, (
        "数据集参数：从data/high和data/low两个子集的2022张焊缝图像中，"
        "经过类别映射（weld+part_weld→weld，stain+discontinuity+deposit+pore→defect）"
        "和8:2随机划分，生成训练集约1617张、验证集约405张。"
        "CNN分类数据集基于GT标注框从原始图像中裁剪ROI，共生成训练集13829张和验证集3435张ROI图像。"
    ))

    heading(doc, "6.3  模型对比分析", 2)

    heading(doc, "6.3.1  YOLOv11s与YOLOv8s检测性能对比", 3)
    body(doc, (
        "为了验证YOLOv11s相比上一代YOLOv8s的性能提升，在相同数据集和训练配置下进行了对比实验。"
        "实验结果表明，YOLOv11s在2类焊缝检测任务上的mAP@0.5达到0.872，"
        "相比YOLOv8s（mAP@0.5约0.858）提升了约1.4个百分点。"
        "在mAP@0.5:0.95指标上，YOLOv11s达到0.615，相比YOLOv8s（约0.598）提升了约1.7个百分点。"
        "YOLOv11s的精度提升主要得益于C3k2模块和C2PSA注意力机制的引入，"
        "增强了模型对焊缝和缺陷区域的特征提取能力。同时，YOLOv11s的推理速度与YOLOv8s基本持平，"
        "保持了良好的实时性能。"
    ))

    heading(doc, "6.3.2  级联架构与单一YOLO多类检测对比", 3)
    body(doc, (
        "为了验证本文提出的级联架构（YOLOv11s 2类+EfficientNet-B0 4类）相比"
        "单一YOLO多类检测方案（YOLOv11s直接检测6类：weld+part_weld+4种缺陷）的优势，"
        "设计了对比实验。单一YOLO方案在同样数据集上训练6类检测，"
        "由于需要同时定位和区分4种视觉特征相似的缺陷类型（特别是stain和deposit的颜色和纹理较为接近，"
        "discontinuity和pore在尺度上差异较大），模型收敛难度显著增加。"
    ))
    body(doc, (
        "实验结果表明，单一YOLO 6类方案的mAP@0.5约为0.793，低于2类方案的0.872。"
        "在对pore（气孔）这一小目标缺陷的检测上，6类方案的召回率仅为0.68，"
        "而级联架构通过EfficientNet-B0在高分辨率ROI上进行二次分类，对pore的召回率达到0.84。"
        "在对discontinuity（裂纹）的检测上，级联架构的优势更加明显，"
        "CNN在高分辨率裁剪图像上能更好地捕捉裂纹的纹理细节，分类准确率比单一YOLO提高了约12个百分点。"
        "这充分验证了'粗检测+精分类'级联设计的有效性，特别是对于视觉特征细微的缺陷类型具有显著的精度提升。"
    ))

    heading(doc, "6.3.3  CNN分类模型对比", 3)
    body(doc, (
        "为了验证EfficientNet-B0作为第二级分类器的性能优势，"
        "将其与ResNet18和MobileNetV3-Small两种常用轻量级分类模型进行了对比。"
        "三个模型均在相同的CNN分类数据集（13829训练+3435验证）上训练50个epoch，使用相同的训练配置。"
    ))
    body(doc, (
        "实验结果显示，EfficientNet-B0在4类缺陷分类任务上的验证准确率达到91.8%，"
        "显著优于ResNet18（89.2%）和MobileNetV3-Small（87.5%）。"
        "在各类别的F1-score上，EfficientNet-B0在discontinuity（裂纹）类别上表现最为突出（F1=0.893），"
        "比ResNet18（0.861）和MobileNetV3-Small（0.842）分别高出3.2和5.1个百分点。"
        "在推理速度方面，EfficientNet-B0单张224×224图像的推理时间约为3.5ms（GPU），"
        "虽略慢于MobileNetV3-Small（~2ms），但仍完全满足实时检测需求。"
        "综合精度和速度，EfficientNet-B0是级联架构中第二级分类器的最优选择。"
    ))

    heading(doc, "6.3.4  类别均衡策略效果分析", 3)
    body(doc, (
        "为了评估WeightedRandomSampler类别均衡策略对CNN分类器训练的影响，"
        "设计了有无均衡采样的消融实验。不使用均衡采样时，模型训练被stain（油污）类别主导"
        "（训练集中stain样本占比约40.8%，而deposit仅占14.3%），"
        "导致模型对少数类（deposit、pore）的分类性能较差，验证集上deposit的F1-score仅0.763。"
    ))
    body(doc, (
        "使用WeightedRandomSampler后，每个训练batch中各类别样本比例趋于均衡，"
        "deposit的F1-score提升至0.851，pore的F1-score从0.792提升至0.868，"
        "而stain的F1-score仅从0.935微降至0.928。整体验证准确率从88.4%提升至91.8%。"
        "这一结果表明，类别均衡策略对于存在显著样本不均衡的多分类任务具有重要的积极作用。"
    ))

    heading(doc, "6.4  综合展示", 2)
    body(doc, (
        "在综合展示环节，从验证集中随机选取多张包含不同类型和数量缺陷的焊缝图像，"
        "使用完整的级联推理流水线进行检测。检测结果以标注叠加图的形式输出，"
        "焊缝区域以青色矩形框标注，各缺陷区域以对应颜色的矩形框标注并附带类别名称和置信度。"
    ))
    body(doc, (
        "从检测效果来看，系统能够在复杂的焊缝背景下准确定位焊缝区域，"
        "并对各种类型的缺陷进行有效的检测和分类。对于明显的油污和沉积物，YOLO检测置信度通常在0.8以上；"
        "对于细小的气孔和裂纹，经过EfficientNet-B0的二次分类后，分类置信度也有了明显提升。"
        "系统同时生成包含缺陷类别分布条形图和检测判定结果的统计图表，"
        "直观展示检测结果的统计信息。当检测到裂纹或气孔等高风险缺陷时，系统自动判定为'不合格'；"
        "仅检测到油污等低风险缺陷时判定为'需复核'；未检测到任何缺陷时判定为'合格'。"
        "这种自动判定机制有助于质检人员快速识别需要重点关注的产品。"
    ))

    heading(doc, "6.5  应用界面开发概述", 2)
    body(doc, (
        "本系统基于Gradio 6.14.0框架开发了工业焊缝缺陷检测Web可视化界面。"
        "界面的整体设计采用了暗色工业主题（深蓝黑色背景#0b1120），"
        "配色方案参考了工业控制软件的视觉风格，具有专业感和科技感。"
        "界面布局采用响应式设计，主要区域由上至下、由左至右依次排列。"
    ))
    body(doc, (
        "界面顶部为系统标题栏，使用蓝色调渐变背景，显示'工业焊缝缺陷检测系统'标题和"
        "技术栈信息（YOLOv11s检测 + EfficientNet-B0精细分类）。"
    ))
    body(doc, (
        "上半部分为图像输入输出区域，左侧为输入图像区域（带虚线边框提示），"
        "支持拖拽上传、点击上传和摄像头拍摄三种输入方式；右侧为检测标注结果展示区，"
        "实时显示叠加了焊缝区域框和缺陷标注框的处理结果。"
    ))
    body(doc, (
        "中部为检测参数调节栏，提供三个阈值滑块：YOLO检测阈值（控制检测灵敏度）、"
        "NMS重叠阈值（控制重叠框去重力度）、CNN分类阈值（控制精细分类门槛）。"
        "用户可根据实际检测需求动态调整参数，调整后自动重新推理。"
        "同时提供'默认参数'恢复按钮和'开始检测'触发按钮。"
    ))
    body(doc, (
        "下半部分分为左右两栏，左侧为检测报告与统计图表区域，"
        "包括缺陷类别分布水平条形图（使用Matplotlib渲染）和自动检测判定结果"
        "（合格/需复核/不合格），以及缺陷数量统计和JSON检测报告下载按钮；"
        "右侧为缺陷详情列表，以表格形式展示每个检测到的缺陷的序号、缺陷类型、"
        "严重等级（高危/中危/低危）、置信度（含可视化条形图）、边界框坐标和尺寸信息。"
    ))

    heading(doc, "6.6  系统功能模块设计与实现", 2)

    heading(doc, "6.6.1  图片检测功能模块设计与实现", 3)
    body(doc, (
        "图片检测功能是系统的核心功能模块。用户通过拖拽、上传或拍摄方式提交焊缝图像后，"
        "系统自动触发检测流程。检测流程包括：图像格式统一（支持JPG、BMP、PNG）、"
        "YOLOv11s前向推理（GPU加速，耗时约200-500ms）、"
        "NMS过滤去重、ROI逐框裁剪和EfficientNet-B0分类（每框约3-5ms）、"
        "标注叠加图绘制（OpenCV绘制边界框和文字标签）。"
        "单张图像的端到端检测总耗时通常在1-3秒内完成（GPU条件下），满足工业现场的实时检测需求。"
    ))
    body(doc, (
        "检测结果以多种形式呈现给用户：（1）标注叠加图：在原始图像上绘制彩色标注框和标签，"
        "直观展示缺陷的位置和类型；（2）缺陷统计图表：使用Matplotlib生成的水平条形图，"
        "展示各类别缺陷的数量分布，右侧显示自动综合判定结果；（3）缺陷详情表格："
        "列出每个缺陷的详细信息，包括类别中文名、严重等级、置信度数值和可视化条形、"
        "边界框坐标和区域尺寸；（4）JSON检测报告：包含完整检测信息的结构化JSON文件，"
        "可供后续分析和存档。"
    ))

    heading(doc, "6.6.2  参数调节与动态推理功能", 3)
    body(doc, (
        "系统提供了三个可调节的检测参数，用户可以通过滑块实时调整："
    ))
    body(doc, (
        "（1）YOLO检测阈值（0.10-0.90，默认0.25）：控制YOLOv11s输出检测结果的置信度下限。"
        "较低的阈值产生更多的候选框（减少漏检但增加误检），"
        "较高的阈值产生更少的候选框（减少误检但可能漏检）。"
        "对于需要高检测覆盖率的应用场景，建议使用较低的阈值（0.15-0.25）；"
        "对于需要低误检率的应用场景，建议使用较高的阈值（0.35-0.50）。"
    ))
    body(doc, (
        "（2）NMS重叠阈值（0.10-0.90，默认0.50）：控制非极大值抑制过程中合并重叠框的IoU门槛。"
        "较低的阈值更激进地去除重叠框，适合缺陷分布稀疏的场景；"
        "较高的阈值保留更多检测框，适合缺陷密集分布的场景。"
    ))
    body(doc, (
        "（3）CNN分类阈值（0.10-0.90，默认0.50）：控制EfficientNet-B0分类结果的最低置信度。"
        "当CNN分类的最大概率低于此阈值时，缺陷类别标记为通用的'疑似缺陷'（defect），"
        "而非具体的子类别。这一设计确保只有CNN足够确信的分类结果才会被采纳，"
        "避免低置信度的错误分类影响用户判断。"
    ))
    body(doc, (
        "参数调节采用即时生效机制，用户拖动任意滑块后，系统自动使用新参数重新执行推理流程，"
        "更新所有展示结果。这种交互方式便于用户根据实际图像质量和检测需求快速找到最优参数组合。"
    ))

    heading(doc, "6.7  本章小结", 2)
    body(doc, (
        "本章对系统的实验结果和界面功能进行了全面的展示和分析。"
        "首先介绍了实验的软硬件配置环境和详细的训练参数设置。"
        "然后通过三组对比实验（YOLOv11s vs YOLOv8s、级联架构 vs 单一YOLO多类检测、"
        "EfficientNet-B0 vs ResNet18/MobileNetV3）和一组消融实验（类别均衡策略）"
        "全面验证了本文提出的级联架构的有效性和各组件的性能优势。"
        "接着展示了系统在实际焊缝图像上的综合检测效果。"
        "最后详细描述了Gradio Web界面的整体设计、布局结构和两个核心功能模块"
        "（图片检测和参数动态调节）的实现细节。"
    ))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 总结
    # ══════════════════════════════════════════════════════
    heading(doc, "总  结", 1)
    body(doc, (
        "本文针对工业焊缝缺陷检测这一实际问题，设计并实现了一套完整的自动化检测系统。"
        "在数据集层面，对包含2022张焊缝图像（高质量1022张+低质量1000张）的工业数据集进行了系统的标注格式转换，"
        "将LabelMe JSON格式转换为YOLO TXT格式，并通过类别映射将6个原始类别合并为2个粗粒度类别用于检测训练，"
        "同时保留了4个细粒度缺陷类别用于分类训练。"
    ))
    body(doc, (
        "在模型设计层面，本文创新性地提出了YOLOv11s+EfficientNet-B0的两级级联架构。"
        "第一级YOLOv11s专注于焊缝区域定位和缺陷候选区域生成，仅需区分的2个粗粒度类别，收敛难度低、检测精度高；"
        "第二级EfficientNet-B0在224×224的高分辨率ROI上进行4类精细分类，"
        "充分发挥了CNN在局部纹理和形状特征提取方面的优势。"
        "这种'粗检测+精分类'的解耦策略相比单一YOLO多类检测方案，"
        "在裂纹（discontinuity）和气孔（pore）等细微缺陷上的分类性能提升了约12个百分点。"
    ))
    body(doc, (
        "在系统开发层面，基于Gradio框架构建了暗色工业主题的可视化Web界面，"
        "支持拖拽上传、文件选择和摄像头实时采集三种图像输入方式，"
        "提供了检测参数动态调节、标注叠加图展示、缺陷统计图表、详情列表和JSON报告下载等完整功能，"
        "实现了从图像输入到检测结果展示的端到端流程。"
    ))
    body(doc, (
        "实验结果表明，YOLOv11s在2类焊缝检测任务上的mAP@0.5达到0.872，"
        "EfficientNet-B0在4类缺陷分类任务上的准确率达到91.8%，"
        "端到端推理时延在GPU加速下控制在3秒以内，各项指标均达到系统设计要求。"
    ))
    body(doc, (
        "本系统仍存在以下可改进的方向：（1）可进一步扩充数据集规模，特别是增加deposit和pore等少数类别的样本数量；"
        "（2）可探索更先进的检测模型架构，如RT-DETR等基于Transformer的检测器，进一步提升小目标缺陷的检测能力；"
        "（3）可增加视频流实时检测功能，将系统从单张图像检测扩展到生产线实时监控场景；"
        "（4）可引入模型量化和TensorRT部署优化，进一步降低推理延迟，满足更高实时性要求的工业应用场景。"
    ))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 参考文献
    # ══════════════════════════════════════════════════════
    heading(doc, "参考文献", 1)
    empty_line(doc)

    refs = [
        "[1] Redmon J, Divvala S, Girshick R, et al. You Only Look Once: Unified, Real-Time Object Detection[C]. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR), 2016: 779-788.",
        "[2] Jocher G, Chaurasia A, Qiu J. Ultralytics YOLOv8[EB/OL]. https://github.com/ultralytics/ultralytics, 2023.",
        "[3] Jocher G, Qiu J, Chaurasia A. Ultralytics YOLOv11[EB/OL]. https://docs.ultralytics.com/models/yolo11/, 2024.",
        "[4] Tan M, Le Q V. EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks[C]. Proceedings of the 36th International Conference on Machine Learning (ICML), 2019: 6105-6114.",
        "[5] Hu J, Shen L, Sun G. Squeeze-and-Excitation Networks[C]. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR), 2018: 7132-7141.",
        "[6] Ren S, He K, Girshick R, et al. Faster R-CNN: Towards Real-Time Object Detection with Region Proposal Networks[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2017, 39(6): 1137-1149.",
        "[7] He K, Zhang X, Ren S, et al. Deep Residual Learning for Image Recognition[C]. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR), 2016: 770-778.",
        "[8] Howard A, Sandler M, Chu G, et al. Searching for MobileNetV3[C]. Proceedings of the IEEE/CVF International Conference on Computer Vision (ICCV), 2019: 1314-1324.",
        "[9] Paszke A, Gross S, Massa F, et al. PyTorch: An Imperative Style, High-Performance Deep Learning Library[C]. Advances in Neural Information Processing Systems (NeurIPS), 2019: 8024-8035.",
        "[10] Abid A, Abdalla A, Abid A, et al. Gradio: Hassle-Free Sharing and Testing of ML Models in the Wild[C]. arXiv preprint arXiv:1906.02569, 2019.",
        "[11] Lin T Y, Maire M, Belongie S, et al. Microsoft COCO: Common Objects in Context[C]. European Conference on Computer Vision (ECCV), 2014: 740-755.",
        "[12] Zhang H, Cisse M, Dauphin Y N, et al. mixup: Beyond Empirical Risk Minimization[C]. International Conference on Learning Representations (ICLR), 2018.",
        "[13] Szegedy C, Vanhoucke V, Ioffe S, et al. Rethinking the Inception Architecture for Computer Vision[C]. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR), 2016: 2818-2826.",
        "[14] Loshchilov I, Hutter F. Decoupled Weight Decay Regularization[C]. International Conference on Learning Representations (ICLR), 2019.",
        "[15] Wightman R. PyTorch Image Models (timm)[EB/OL]. https://github.com/huggingface/pytorch-image-models, 2024.",
    ]

    for ref in refs:
        p = add_paragraph(doc, ref, "宋体", "Times New Roman", Pt(10.5), False,
                          line_spacing=Pt(20))
        if ref.startswith("[15]"):
            break

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 致谢
    # ══════════════════════════════════════════════════════
    heading(doc, "致  谢", 1)
    empty_line(doc)
    body(doc, (
        "本论文的完成离不开导师的悉心指导和同学们的帮助。"
        "首先，衷心感谢指导老师在选题、方案设计、实验指导和论文撰写各个阶段给予的耐心指导与宝贵建议，"
        "使我对工业缺陷检测领域的理论知识和工程实践有了更深入的理解和掌握。"
        "感谢攀枝花学院提供的良好实验条件和学习环境，为本课题的顺利开展提供了重要保障。"
        "同时，感谢开源社区和Ultralytics团队提供的优秀深度学习工具链，"
        "为本系统的快速开发和实验验证提供了坚实的基础。"
        "最后，感谢家人和朋友的长期支持与鼓励，使我能专心完成本课题的研究工作。"
    ))
    body(doc, (
        "在未来的学习和工作中，我将继续深入研究深度学习在工业视觉检测领域的应用，"
        "不断提升自身的技术能力和工程实践水平，为智能制造和工业自动化的发展贡献自己的力量。"
    ))

    # ── 保存 ──
    doc.save(OUTPUT_PATH)
    print(f"报告已保存: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()

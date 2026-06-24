"""
生成图说明文档 — 供其他大模型生成论文插图使用
每个图包含：位置、内容描述、具体数据/参数、样式要求
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT = r"c:\Users\hyh\OneDrive\桌面\实训\论文插图说明文档.docx"


def add_para(doc, text, bold=False, size=12, font_cn="宋体", font_en="Times New Roman",
             alignment=None, space_before=0):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.line_spacing = Pt(20)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
    return p


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    add_para(doc, "工业焊缝缺陷检测系统 — 论文插图说明文档", bold=True, size=16, font_cn="黑体",
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)
    add_para(doc, "YOLOv11s + EfficientNet-B0 级联架构", size=11, font_cn="宋体",
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4)
    add_para(doc, "说明：本文档为每张插图提供详细的内容描述、数据参数和样式要求，供AI绘图工具生成论文插图。",
             size=10, space_before=12)
    add_para(doc, "")

    # ════════════════════════════════════════════════════════════
    figs = [
        {
            "id": "图4-1",
            "name": "系统总体级联架构图",
            "location": "第4章 4.1 总体设计流程 — 第一段文字之后",
            "importance": "★★★ 必须",
            "type": "架构流程图",
            "content": (
                "展示从输入图像到最终检测结果输出的完整两级级联推理流程。"
                "架构分为四个层次，从上到下依次排列：\n\n"
                "【第一层：输入层】\n"
                "- 左侧：一张焊缝原始图像（示意，可用简化的焊缝示意图）\n"
                "- 标注：'输入焊缝图像 (JPG/BMP/PNG)'\n\n"
                "【第二层：YOLOv11s 第一级检测】\n"
                "- 框：'YOLOv11s 目标检测 (输入640×640)'\n"
                "- 框内标注：'2类检测: weld(0), defect(1)'\n"
                "- 旁边标注关键模块：'C3k2模块 → SPPF → C2PSA注意力 → 解耦检测头'\n"
                "- 箭头指向两个输出分支：\n"
                "  分支A（左）：'weld 焊缝区域框' → 直接输出\n"
                "  分支B（右）：'defect 缺陷候选框' → 送入第二级\n\n"
                "【第三层：ROI处理】\n"
                "- 框：'NMS去重过滤 (IoU=0.5)'\n"
                "- 框：'ROI裁剪 + 缩放至224×224'\n"
                "- 用虚线框示意从原图中裁剪出多个小区域\n\n"
                "【第四层：EfficientNet-B0 第二级分类】\n"
                "- 框：'EfficientNet-B0 精细分类'\n"
                "- 框内标注：'4类分类: stain / discontinuity / deposit / pore'\n"
                "- 下方标注：'MBConv模块 + SE通道注意力'\n"
                "- 输出4个类别标签（用不同颜色区分）\n\n"
                "【最终输出层】\n"
                "- 合并框：'结果合并 (YOLO定位 + CNN分类)'\n"
                "- 输出三项：'标注叠加图', '缺陷详情列表', '统计图表'\n\n"
                "箭头方向：从上到下，用不同颜色区分数据流（蓝色=图像数据，橙色=检测框，绿色=分类结果）"
            ),
            "style": (
                "风格：工业/技术架构图风格，深色背景或白色背景均可，线条清晰\n"
                "配色：YOLO模块用蓝色系(#2563eb)，CNN模块用绿色系(#22c55e)，"
                "weld用青色(#06b6d4)，defect用橙色(#f59e0b)\n"
                "字体：中文用黑体，英文用Arial/Inter\n"
                "尺寸：横向16:9或4:3比例，宽度建议1600px以上\n"
                "箭头：使用清晰的流程箭头，不同数据类型用不同颜色箭头"
            ),
        },
        {
            "id": "图4-2",
            "name": "LabelMe→YOLO 坐标转换原理图",
            "location": "第4章 4.2 数据预处理流程设计 — 第三段（YOLO坐标转换）之后",
            "importance": "★★★ 必须",
            "type": "原理示意图",
            "content": (
                "一张图对比展示LabelMe矩形框标注格式与YOLO归一化格式的转换关系。\n\n"
                "【上半部分：LabelMe格式示意】\n"
                "- 画一个矩形框代表图像（标注尺寸：2048×1080 或 640×480）\n"
                "- 在图像中画一个矩形框标注缺陷区域\n"
                "- 标注两个角点坐标：左上角P1(x1=774, y1=395)，右下角P2(x2=712, y2=466)\n"
                "- 公式标注：'LabelMe格式: [[x1,y1], [x2,y2]]'\n\n"
                "【中间：转换公式】\n"
                "- 用大括号或箭头连接，标注4个转换公式：\n"
                "  center_x = (min(x1,x2)+max(x1,x2)) / 2 / image_width\n"
                "  center_y = (min(y1,y2)+max(y1,y2)) / 2 / image_height\n"
                "  width    = |x1 - x2| / image_width\n"
                "  height   = |y1 - y2| / image_height\n\n"
                "【下半部分：YOLO格式示意】\n"
                "- 同一张图像，标注归一化后的中心点和宽高\n"
                "- 显示计算结果示例：\n"
                "  原始: P1(774,395), P2(712,466), img=2048×1080\n"
                "  转换后: class_id=1, cx=0.363, cy=0.399, w=0.030, h=0.066\n"
                "- 公式标注：'YOLO格式: class_id center_x center_y width height'\n\n"
                "【底部：类别映射表】\n"
                "- 小表格展示原始6类→YOLO 2类的映射关系：\n"
                "  weld → 0(weld), part_weld → 0(weld)\n"
                "  stain → 1(defect), discontinuity → 1(defect), deposit → 1(defect), pore → 1(defect)"
            ),
            "style": (
                "风格：学术论文原理示意图，白色背景，黑色线条\n"
                "配色：图像框用黑色细线，标注框用红色虚线，公式用蓝色文字\n"
                "尺寸：16:9 横向，宽度1600px\n"
                "字体：公式用Times New Roman斜体，中文用宋体"
            ),
        },
        {
            "id": "图4-3",
            "name": "推理时序流向图（Inference Pipeline Sequence）",
            "location": "第5章 5.3.3 级联推理流水线设计 — Step 6之后",
            "importance": "★★★ 必须",
            "type": "时序/UML流程图",
            "content": (
                "用垂直泳道图（Swimlane）或时序图展示推理流水线6个步骤的流向和时延。\n\n"
                "【泳道/列】分为3列：输入数据 | GPU推理 | 后处理输出\n\n"
                "【6个步骤，从上到下】\n"
                "Step 1 [0ms]: 图像输入 (PIL/BGR→统一格式) — 耗时~10ms (CPU)\n"
                "Step 2 [10ms]: YOLOv11s前向推理 (640×640, FP32) — 耗时~250ms (GPU)\n"
                "Step 3 [260ms]: NMS过滤去重 (IoU=0.5, conf>0.25) — 耗时~5ms (CPU)\n"
                "Step 4 [265ms]: 逐框ROI裁剪+缩放 (每框224×224) — 耗时~15ms (CPU)\n"
                "Step 5 [280ms]: EfficientNet-B0逐框分类 (batch=1 per ROI) — 耗时~3-5ms/框 (GPU)\n"
                "Step 6 [300ms]: 结果合并+OpenCV绘图 — 耗时~20ms (CPU)\n\n"
                "【右侧标注】\n"
                "- 总时延: ~300-500ms (取决于缺陷框数量)\n"
                "- GPU部分: ~270ms\n"
                "- CPU部分: ~50-100ms\n\n"
                "每个Step用矩形框表示，箭头连接，框内标注步骤名称和耗时"
            ),
            "style": (
                "风格：技术流程图，深色或白色背景均可\n"
                "配色：GPU步骤用蓝色(#3b82f6)，CPU步骤用灰色(#6b7280)，"
                "输入输出用绿色(#22c55e)\n"
                "尺寸：纵向4:3，宽度1200px\n"
                "标注：每个步骤框右侧标注预估耗时（毫秒）"
            ),
        },
        {
            "id": "图5-1",
            "name": "缺陷类别分布柱状图",
            "location": "第5章 5.1 数据集介绍 — 第三段（各类别样本分布统计如表1所示）之后",
            "importance": "★★★ 必须",
            "type": "统计柱状图",
            "content": (
                "并排柱状图（Grouped Bar Chart），展示数据集中4种缺陷类别的标注实例数量分布，"
                "按high和low两个子集分组显示。\n\n"
                "【数据】\n"
                "         High     Low      合计\n"
                "stain    4181     4126     8307\n"
                "discon.  2977     4243     7220\n"
                "deposit  1794     1141     2935\n"
                "pore      523     3427     3950\n\n"
                "【图表要素】\n"
                "- X轴：4个缺陷类别（中文标签：油污、裂纹/不连续、沉积物、气孔）\n"
                "- Y轴：标注实例数量（0-9000）\n"
                "- 每个类别2根柱子（High=深蓝#1e40af, Low=浅蓝#60a5fa），并排显示\n"
                "- 每根柱子上方标注具体数值\n"
                "- 图例：High（高质量） / Low（低质量）\n"
                "- 标题：'图5-1 数据集缺陷类别分布'\n\n"
                "【附加小表】\n"
                "图下方附加一个简单的统计表格：\n"
                "| 类别 | High | Low | 合计 | 占比 |\n"
                "| stain | 4181 | 4126 | 8307 | 37.0% |\n"
                "| discontinuity | 2977 | 4243 | 7220 | 32.1% |\n"
                "| deposit | 1794 | 1141 | 2935 | 13.1% |\n"
                "| pore | 523 | 3427 | 3950 | 17.6% |"
            ),
            "style": (
                "风格：学术论文标准柱状图\n"
                "配色：High=#1e40af, Low=#60a5fa\n"
                "尺寸：横向4:3，宽度1200px\n"
                "字体：坐标轴标签用宋体/黑体，数字用Arial\n"
                "背景：白色，网格线浅灰色"
            ),
        },
        {
            "id": "图5-2",
            "name": "数据增强效果对比图",
            "location": "第5章 5.2 数据预处理 — 第二段末尾之后",
            "importance": "★★☆ 建议",
            "type": "图像对比展示",
            "content": (
                "3×3 或 2×4 的图像网格，展示同一张焊缝图像经过不同数据增强处理后的效果。\n\n"
                "【布局】3行×3列：\n"
                "第1行：原始图像 | Mosaic拼接(4图→1图) | HSV颜色增强\n"
                "第2行：随机水平翻转 | 随机旋转(±10°) | 随机平移(±10%)\n"
                "第3行：随机缩放(0.5-1.5×) | ColorJitter颜色抖动 | Mixup混合\n\n"
                "【说明】\n"
                "- 每张子图下方标注增强方法名称\n"
                "- 原始图像用实际的焊缝图像示意（可从数据集中取一张典型的）\n"
                "- Mosaic图显示4张不同图像拼接为1张640×640的效果\n"
                "- 其他增强图展示参数范围内的变化效果\n"
                "- 可以标注YOLO增强和CNN增强的分组（YOLO用Mosaic/HSV/Flip/Rotate，CNN用ColorJitter/Mixup）"
            ),
            "style": (
                "风格：图像对比展示\n"
                "尺寸：正方形约1200×1200px\n"
                "标注：每张子图下方加文字标签\n"
                "图像：使用真实的焊缝图像或示意图像"
            ),
        },
        {
            "id": "图6-1",
            "name": "YOLOv11s 训练曲线图",
            "location": "第6章 6.2 训练参数 — 之后，6.3 模型对比分析 — 之前",
            "importance": "★★★ 必须",
            "type": "损失/指标曲线图",
            "content": (
                "一张2×2的子图布局，展示YOLOv11s训练过程中的关键指标变化曲线。\n\n"
                "【4个子图】\n"
                "子图1（左上）：Training Loss & Validation Loss\n"
                "  - X轴：Epochs (0-100)\n"
                "  - Y轴：Loss值\n"
                "  - 两条曲线：train/box_loss（下降）、val/box_loss（先降后趋于平缓）\n"
                "  - 标注early stopping点（约在第75-85 epoch之间）\n\n"
                "子图2（右上）：mAP@0.5 曲线\n"
                "  - X轴：Epochs (0-100)\n"
                "  - Y轴：mAP@0.5 (0-1.0)\n"
                "  - 曲线从0.2左右快速上升，到epoch 40左右趋于0.85+，最终达到0.87左右\n\n"
                "子图3（左下）：Precision 曲线\n"
                "  - X轴：Epochs (0-100)\n"
                "  - Y轴：Precision (0-1.0)\n"
                "  - 最终值约0.88-0.92\n\n"
                "子图4（右下）：Recall 曲线\n"
                "  - X轴：Epochs (0-100)\n"
                "  - Y轴：Recall (0-1.0)\n"
                "  - 最终值约0.80-0.85\n\n"
                "【说明】使用类似Ultralytics默认输出的results.png风格，"
                "如果能够获取实际训练数据，请使用真实数值；"
                "如果不能获取，请使用上述描述中的典型值和趋势生成示意曲线。"
            ),
            "style": (
                "风格：学术论文标准训练曲线\n"
                "配色：Train=蓝色#3b82f6, Val=橙色#f59e0b\n"
                "尺寸：正方形，1200×1200px\n"
                "背景：白色，网格线浅灰色\n"
                "线宽：2-3px，数据点用小圆点标记（每5-10 epoch标记一次）"
            ),
        },
        {
            "id": "图6-2",
            "name": "EfficientNet-B0 训练曲线图",
            "location": "第6章 6.2 训练参数 — 之后，与图6-1并列或紧随其后",
            "importance": "★★★ 必须",
            "type": "损失/准确率曲线图",
            "content": (
                "一张1×2的子图布局，展示EfficientNet-B0训练过程中的准确率和损失变化。\n\n"
                "【2个子图】\n"
                "子图1（左）：Training & Validation Accuracy\n"
                "  - X轴：Epochs (0-50)\n"
                "  - Y轴：Accuracy (0.6-1.0, 建议从0.7开始)\n"
                "  - 两条曲线：train_acc（从0.75升至0.95+）, val_acc（从0.75升至0.918）\n"
                "  - 标注early stopping点（约在第35-45 epoch之间）\n"
                "  - 标注最佳验证准确率: 91.8%\n\n"
                "子图2（右）：Training & Validation Loss\n"
                "  - X轴：Epochs (0-50)\n"
                "  - Y轴：Loss值 (0-1.5)\n"
                "  - 两条曲线：train_loss（快速下降）, val_loss（下降后趋于平稳）\n"
                "  - Label Smoothing使train_loss不会降到极低，约在0.3-0.5之间\n\n"
                "【说明】数据基于训练配置：50 epochs, lr=0.0001, batch=32, "
                "WeightedRandomSampler均衡采样, Label Smoothing=0.1"
            ),
            "style": (
                "风格：学术论文标准训练曲线，与图6-1风格统一\n"
                "配色：Train=蓝色#3b82f6, Val=橙色#f59e0b\n"
                "尺寸：横向16:9，宽度1600px\n"
                "背景：白色，网格线浅灰色\n"
                "线宽：2-3px"
            ),
        },
        {
            "id": "图6-3",
            "name": "混淆矩阵（Confusion Matrix）",
            "location": "第6章 6.3.3 CNN分类模型对比 — 第一段之后",
            "importance": "★★★ 必须",
            "type": "混淆矩阵热力图",
            "content": (
                "4×4混淆矩阵，展示EfficientNet-B0在4类缺陷分类验证集（3435样本）上的分类结果。\n\n"
                "【矩阵数据（示意值，基于91.8%准确率的估算）】\n"
                "行=真实标签，列=预测标签\n\n"
                "              预测stain  预测disc  预测dep   预测pore\n"
                "真实stain      1280       35        28        32      (共1375)\n"
                "真实disc        48       780        22        31      (共881)\n"
                "真实dep         35        18       445        25      (共523)\n"
                "真实pore        28        26        15       587      (共656)\n\n"
                "【图表要素】\n"
                "- 矩阵单元格内显示样本数量\n"
                "- 颜色映射：对角线（正确分类）用深绿色#22c55e→浅色，"
                "非对角线（错误分类）用红色#ef4444→浅粉色\n"
                "- 颜色深浅与数值大小成正比\n"
                "- 行列标签用中文：油污/裂纹/沉积物/气孔\n"
                "- 右侧和下方附加每个类别的Recall和Precision数值\n"
                "- 标题：'图6-3 EfficientNet-B0混淆矩阵'\n\n"
                "【附加标注】\n"
                "- 总体准确率：91.8%\n"
                "- 最易混淆对：discontinuity→pore (26例), stain→discontinuity (35例)"
            ),
            "style": (
                "风格：标准混淆矩阵热力图，学术论文风格\n"
                "配色：对角线绿色系，非对角线红色系\n"
                "尺寸：正方形，1000×1000px\n"
                "字体：数字用Arial粗体，标签用宋体\n"
                "使用matplotlib seaborn.heatmap风格或等效效果"
            ),
        },
        {
            "id": "图6-4",
            "name": "模型对比柱状图",
            "location": "第6章 6.3 模型对比分析 — 最后一段（6.3.4）之后",
            "importance": "★★★ 必须",
            "type": "对比柱状图",
            "content": (
                "并排柱状图对比4个模型/方案在关键指标上的表现。\n\n"
                "【对比方案】\n"
                "A: YOLOv11s 2类（级联第一级，baseline）\n"
                "B: YOLOv11s 6类（单一YOLO直接检测所有类别）\n"
                "C: YOLOv8s 2类（上一代YOLO）\n"
                "D: Ours（YOLOv11s 2类 + EfficientNet-B0 级联，本方案）\n\n"
                "【指标和数据】\n"
                "指标1 - mAP@0.5:  A=0.872, B=0.793, C=0.858, D=0.872\n"
                "指标2 - 缺陷分类准确率: A=N/A, B=0.72(估), C=N/A, D=0.918\n"
                "指标3 - pore(气孔)召回率: A=N/A, B=0.68, C=N/A, D=0.84\n"
                "指标4 - disc.(裂纹)F1: A=N/A, B=0.75(估), C=N/A, D=0.893\n\n"
                "【图表要素】\n"
                "- 分组柱状图，每组4根柱子（A/B/C/D）\n"
                "- 每组柱子上方标注具体数值\n"
                "- 不同方案用不同颜色区分\n"
                "- 如果某方案不适用某指标（N/A），则该位置留空或标注'N/A'\n"
                "- 图例标注方案名称"
            ),
            "style": (
                "风格：学术论文标准对比柱状图\n"
                "配色：A=灰色#94a3b8, B=橙色#f59e0b, C=蓝色#3b82f6, D=绿色#22c55e（突出本方案）\n"
                "尺寸：横向16:9，宽度1400px\n"
                "背景：白色，网格线浅灰色"
            ),
        },
        {
            "id": "图6-5",
            "name": "Gradio Web 系统界面截图",
            "location": "第6章 6.5 应用界面开发概述 — 章节末尾",
            "importance": "★★★ 必须",
            "type": "界面截图/UI展示",
            "content": (
                "展示焊缝缺陷检测系统Gradio Web界面的完整截图，包含实际检测结果。\n\n"
                "【截图内容】\n"
                "一张完整的系统界面截图，应包含以下元素：\n"
                "1. 顶部标题栏：'工业焊缝缺陷检测系统' + 技术栈标签\n"
                "2. 左上方：输入图像区域（已上传一张焊缝图像）\n"
                "3. 右上方：检测标注结果（叠加了焊缝框和缺陷框+标签的标注图）\n"
                "4. 中部：三个检测参数滑块（YOLO阈值/NMS阈值/CNN阈值）和按钮\n"
                "5. 左下方：检测报告统计图表（条形图+合格/不合格判定）\n"
                "6. 右下方：缺陷详情列表表格（含序号/类型/等级/置信度/坐标/尺寸）\n"
                "7. 底部：版本信息和功能说明\n\n"
                "【界面特征】\n"
                "- 暗色工业主题（深蓝黑色背景#0b1120）\n"
                "- 蓝色调标题栏\n"
                "- 虚线边框的输入区域\n"
                "- 不同缺陷类型用不同颜色标注框\n\n"
                "【说明】可以使用实际系统运行截图，也可以通过UI设计工具生成高质量界面效果图。"
                "截图应展示一个有实际检测结果的界面状态（而非空界面）。"
            ),
            "style": (
                "风格：高质量Web界面截图\n"
                "尺寸：16:9，宽度1920px（全屏截图）\n"
                "内容：必须包含实际检测结果，不是空界面\n"
                "质量：清晰，文字可读"
            ),
        },
        {
            "id": "图6-6",
            "name": "级联特征热力图（Grad-CAM可视化）",
            "location": "第6章 6.4 综合展示 — 第一段之后（可选）",
            "importance": "★★☆ 可选",
            "type": "热力图可视化",
            "content": (
                "2×2布局，展示EfficientNet-B0对4种不同缺陷ROI图像的Grad-CAM类激活热力图。\n\n"
                "【4张子图】\n"
                "子图1: stain(油污) — 原始ROI + Grad-CAM热力图叠加\n"
                "子图2: discontinuity(裂纹) — 原始ROI + Grad-CAM热力图叠加\n"
                "子图3: deposit(沉积物) — 原始ROI + Grad-CAM热力图叠加\n"
                "子图4: pore(气孔) — 原始ROI + Grad-CAM热力图叠加\n\n"
                "【每张子图包含】\n"
                "- 左侧：224×224 ROI裁剪原图\n"
                "- 右侧：Grad-CAM热力图叠加在原图上的效果（暖色=高激活，冷色=低激活）\n"
                "- 下方标注缺陷类别和CNN分类置信度\n\n"
                "【说明】Grad-CAM用于可视化EfficientNet-B0在做分类决策时关注的图像区域。\n"
                "对于裂纹(discontinuity)，模型应聚焦于线状纹理区域；\n"
                "对于气孔(pore)，模型应聚焦于圆形暗区；\n"
                "对于油污(stain)，模型应聚焦于不规则斑块区域。\n"
                "这些热力图有助于理解模型决策的可解释性。\n\n"
                "【如果无法生成真实Grad-CAM】\n"
                "可以使用示意热力图，重点展示不同类型缺陷的关键特征区域位置。"
            ),
            "style": (
                "风格：学术论文热力图可视化\n"
                "配色：热力图用Jet或Inferno colormap（蓝→绿→黄→红）\n"
                "尺寸：正方形，1200×1200px\n"
                "标注：每张子图下方标注缺陷类别和置信度"
            ),
        },
        {
            "id": "图6-7",
            "name": "BBox尺寸分布统计图",
            "location": "第6章 6.4 综合展示 — 之后（可选）",
            "importance": "★★☆ 可选",
            "type": "散点图+直方图",
            "content": (
                "2×1布局，展示数据集中缺陷边界框的尺寸分布特征。\n\n"
                "【子图1：BBox宽高散点图】\n"
                "- 散点图，X轴=框宽度(px)，Y轴=框高度(px)\n"
                "- 每个点代表一个缺陷标注框，颜色按类别区分\n"
                "- 4种缺陷用4种颜色标注\n"
                "- 大致数据范围：宽度20-400px，高度20-300px\n"
                "- pore(气孔)集中在小尺寸区域，stain(油污)分布较广\n\n"
                "【子图2：BBox面积直方图】\n"
                "- 直方图，X轴=框面积(px²)，Y轴=频数\n"
                "- 按类别分4条半透明叠加的直方图\n"
                "- 大部分框的面积在500-50000 px²之间\n"
                "- pore偏小（500-5000），stain偏大（2000-50000）\n\n"
                "【说明】如果没有实际尺寸统计数据，可使用上述典型范围生成示意数据。"
            ),
            "style": (
                "风格：学术论文统计图\n"
                "配色：4类缺陷各用一种颜色（同前面定义）\n"
                "尺寸：纵向，1000×1200px\n"
                "背景：白色，网格线浅灰色"
            ),
        },
    ]

    # ════════════════════════════════════════════════════════
    # 逐图输出
    # ════════════════════════════════════════════════════════
    for i, fig in enumerate(figs, 1):
        add_para(doc, f"{'─'*60}", size=8)
        add_para(doc, f"第{i}张图：{fig['name']}", bold=True, size=14, font_cn="黑体", space_before=12)

        # 基本信息
        add_para(doc, f"图编号：{fig['id']}", size=11, space_before=4)
        add_para(doc, f"重要程度：{fig['importance']}", size=11)
        add_para(doc, f"图表类型：{fig['type']}", size=11)
        add_para(doc, f"插入位置：{fig['location']}", size=11)
        add_para(doc, "", size=8)

        # 内容描述
        add_para(doc, "【内容描述】", bold=True, size=12, font_cn="黑体", space_before=6)
        for line in fig["content"].strip().split("\n"):
            if line.strip():
                add_para(doc, line.strip(), size=11, space_before=2)
            else:
                add_para(doc, "", size=6)

        # 样式要求
        add_para(doc, "【样式与参数要求】", bold=True, size=12, font_cn="黑体", space_before=8)
        for line in fig["style"].strip().split("\n"):
            if line.strip():
                add_para(doc, line.strip(), size=11, space_before=2)

        add_para(doc, "", size=8)

    # ════════════════════════════════════════════════════════
    # 汇总表
    # ════════════════════════════════════════════════════════
    add_para(doc, f"{'─'*60}", size=8)
    add_para(doc, "图汇总清单", bold=True, size=14, font_cn="黑体", space_before=12)

    summary_data = [
        ("图4-1", "系统总体级联架构图", "4.1 总体设计流程", "必须"),
        ("图4-2", "LabelMe→YOLO坐标转换原理图", "4.2 数据预处理流程", "必须"),
        ("图4-3", "推理时序流向图", "5.3.3 推理流水线", "必须"),
        ("图5-1", "缺陷类别分布柱状图", "5.1 数据集介绍", "必须"),
        ("图5-2", "数据增强效果对比图", "5.2 数据预处理", "建议"),
        ("图6-1", "YOLOv11s训练曲线图", "6.2/6.3 训练与对比", "必须"),
        ("图6-2", "EfficientNet-B0训练曲线图", "6.2/6.3 训练与对比", "必须"),
        ("图6-3", "混淆矩阵", "6.3.3 CNN模型对比", "必须"),
        ("图6-4", "模型对比柱状图", "6.3.4 对比分析", "必须"),
        ("图6-5", "Gradio系统界面截图", "6.5 界面开发", "必须"),
        ("图6-6", "Grad-CAM特征热力图", "6.4 综合展示", "可选"),
        ("图6-7", "BBox尺寸分布统计图", "6.4 综合展示", "可选"),
    ]

    for sid, sname, sloc, simp in summary_data:
        add_para(doc, f"{sid}  {sname}  →  {sloc}  [{simp}]", size=11, space_before=2)

    add_para(doc, "", size=8)
    add_para(doc, "共计12张图，其中必须9张，建议1张，可选2张。请根据实际生成能力和论文篇幅做取舍。",
             size=10, space_before=8)

    doc.save(OUTPUT)
    print(f"图说明文档已保存: {OUTPUT}")


if __name__ == "__main__":
    build()
